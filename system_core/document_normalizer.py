#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Apply safe DOCX normalization fixes from existing audit JSON logs."""

from __future__ import annotations

import argparse
import json
import shutil
import sys
import time
from pathlib import Path
from typing import Any

THIS_DIR = Path(__file__).resolve().parent
if str(THIS_DIR) not in sys.path:
    sys.path.insert(0, str(THIS_DIR))

from document_model import block_paragraph_refs_docx, default_project_paths, ensure_project_dirs, source_relative


SAFE_FIX_MODES = {"safe_replace", "safe", "replace", "auto_replace"}
HIGH_CONFIDENCE = {"high", "высокая", "высокий", "safe"}


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig", errors="replace"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def iter_audit_logs(logs_dir: Path) -> list[Path]:
    return sorted(path for path in logs_dir.rglob("*__audit.json") if path.is_file())


def source_from_audit(input_dir: Path, audit: dict[str, Any], audit_path: Path, logs_dir: Path) -> Path:
    if input_dir.is_file():
        return input_dir
    rel = str(audit.get("source_relative_path") or "").strip()
    if rel:
        return input_dir / rel
    fallback = audit_path.relative_to(logs_dir).with_name(audit_path.name.replace("__audit.json", ".docx"))
    return input_dir / fallback


def is_high_confidence(value: Any) -> bool:
    text = str(value or "").strip().lower().replace("_", " ").replace("-", " ")
    return any(token in text for token in HIGH_CONFIDENCE)


def issue_to_candidate(issue: dict[str, Any], source_rel: str) -> tuple[dict[str, Any] | None, dict[str, Any]]:
    base = {
        "issue_id": issue.get("issue_id", ""),
        "source_relative_path": source_rel,
        "block_id": issue.get("block_id", ""),
        "location": issue.get("human_location") or issue.get("location") or "",
        "quote": issue.get("quote", ""),
        "problem": issue.get("problem") or issue.get("violation") or "",
        "recommendation": issue.get("recommendation") or issue.get("fix") or "",
        "old_text": issue.get("old_text", ""),
        "new_text": issue.get("new_text", ""),
        "fix_mode": issue.get("fix_mode", ""),
        "confidence": issue.get("confidence", ""),
        "status": "unresolved",
        "status_note": "",
    }
    mode = str(base["fix_mode"] or "").strip().lower()
    if mode not in SAFE_FIX_MODES:
        base["status_note"] = f"fix_mode is {mode or 'requires_review'}"
        return None, base
    if not is_high_confidence(base["confidence"]):
        base["status_note"] = f"confidence is {base['confidence'] or 'unknown'}"
        return None, base
    if not str(base["block_id"] or "").strip():
        base["status_note"] = "block_id missing"
        return None, base
    if not str(base["old_text"] or "").strip() or not str(base["new_text"] or "").strip():
        base["status_note"] = "old_text/new_text missing"
        return None, base
    return dict(base), base


def normalized_output_path(output_dir: Path, input_dir: Path, source: Path) -> Path:
    rel = source_relative(source, input_dir)
    return output_dir / rel.parent / f"{source.stem}__normalized{source.suffix.lower()}"


def apply_candidates(source: Path, out_path: Path, candidates: list[dict[str, Any]], *, dry_run: bool = False) -> dict[str, Any]:
    if source.suffix.lower() != ".docx":
        for item in candidates:
            item["status"] = "unresolved"
            item["status_note"] = "normalization is implemented for DOCX only"
        return {"status": "skipped", "applied": 0, "unresolved": len(candidates), "normalized": ""}

    if not dry_run:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, out_path)
        target = out_path
    else:
        target = source

    refs_payload = block_paragraph_refs_docx(target)
    doc = refs_payload["document"]
    refs = refs_payload["refs"]
    applied = 0

    for item in candidates:
        block_id = str(item.get("block_id") or "").strip()
        old_text = str(item.get("old_text") or "")
        new_text = str(item.get("new_text") or "")
        paragraph = refs.get(block_id)
        if paragraph is None:
            item["status"] = "unresolved"
            item["status_note"] = "block_id not found in DOCX"
            continue
        current = paragraph.text
        count = current.count(old_text)
        if count != 1:
            item["status"] = "unresolved"
            item["status_note"] = f"old_text occurrence count is {count}"
            continue
        if not dry_run:
            paragraph.text = current.replace(old_text, new_text, 1)
        item["status"] = "applied"
        item["status_note"] = ""
        applied += 1

    if not dry_run:
        doc.save(str(out_path))
    unresolved = sum(1 for item in candidates if item.get("status") != "applied")
    return {
        "status": "ok",
        "applied": applied,
        "unresolved": unresolved,
        "normalized": "" if dry_run else str(out_path),
    }


def write_report_docx(out_path: Path, payload: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Tahoma"
    style.font.size = Pt(10)

    doc.add_heading("Audion Docs AI - Document Normalization Report", level=0)
    doc.add_paragraph(f"Provider: {payload.get('provider', '')}")
    doc.add_paragraph(f"Documents: {payload.get('documents', 0)}")
    doc.add_paragraph(f"Applied fixes: {len(payload.get('applied_items', []) or [])}")
    doc.add_paragraph(f"Unresolved items: {len(payload.get('unresolved_items', []) or [])}")

    def fill_table(title: str, rows: list[dict[str, Any]], headers: list[str], getters: list[str]) -> None:
        if not rows:
            return
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for row in rows:
            cells = table.add_row().cells
            for index, key in enumerate(getters):
                cells[index].text = str(row.get(key, ""))
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    fill_table(
        "Applied Fixes",
        payload.get("applied_items", []) or [],
        ["File", "Block", "Old text", "New text", "Reason"],
        ["source_relative_path", "block_id", "old_text", "new_text", "recommendation"],
    )
    fill_table(
        "Unresolved Items",
        payload.get("unresolved_items", []) or [],
        ["File", "Block", "Quote", "Problem", "Recommendation", "Reason"],
        ["source_relative_path", "block_id", "quote", "problem", "recommendation", "status_note"],
    )
    doc.save(str(out_path))


def write_report_xlsx(out_path: Path, payload: dict[str, Any]) -> None:
    from openpyxl import Workbook

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Applied"
    ws.append(["File", "Block", "Old text", "New text", "Reason"])
    for item in payload.get("applied_items", []) or []:
        ws.append([item.get("source_relative_path", ""), item.get("block_id", ""), item.get("old_text", ""), item.get("new_text", ""), item.get("recommendation", "")])
    ws2 = wb.create_sheet("Unresolved")
    ws2.append(["File", "Block", "Quote", "Problem", "Recommendation", "Reason"])
    for item in payload.get("unresolved_items", []) or []:
        ws2.append([item.get("source_relative_path", ""), item.get("block_id", ""), item.get("quote", ""), item.get("problem", ""), item.get("recommendation", ""), item.get("status_note", "")])
    wb.save(str(out_path))


def write_report_md(out_path: Path, payload: dict[str, Any]) -> None:
    lines = [
        "# Audion Docs AI - Document Normalization Report",
        "",
        f"- Provider: {payload.get('provider', '')}",
        f"- Documents: {payload.get('documents', 0)}",
        f"- Applied fixes: {len(payload.get('applied_items', []) or [])}",
        f"- Unresolved items: {len(payload.get('unresolved_items', []) or [])}",
        "",
    ]
    if payload.get("unresolved_items"):
        lines.extend(["## Unresolved Items", ""])
        for item in payload["unresolved_items"]:
            lines.append(f"- `{item.get('source_relative_path', '')}` `{item.get('block_id', '')}`: {item.get('status_note', '')}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")


def normalize_from_audit_logs(
    *,
    root: Path,
    provider: str,
    logs_dir: Path,
    input_dir: Path,
    output_dir: Path,
    report_dir: Path,
    patch_dir: Path,
    dry_run: bool = False,
) -> dict[str, Any]:
    paths = default_project_paths(root)
    ensure_project_dirs(paths)
    provider = str(provider or "").strip().lower()
    applied_items: list[dict[str, Any]] = []
    unresolved_items: list[dict[str, Any]] = []
    document_results: list[dict[str, Any]] = []

    for audit_path in iter_audit_logs(logs_dir):
        audit = read_json(audit_path)
        audit_provider = str((audit.get("meta") or {}).get("provider") or "").strip().lower()
        if provider and audit_provider and audit_provider != provider:
            continue
        if provider and not audit_provider:
            continue
        source = source_from_audit(input_dir, audit, audit_path, logs_dir)
        try:
            source_rel = source_relative(source, input_dir).as_posix() if source.exists() else ""
        except ValueError:
            source_rel = ""
        source_rel = source_rel or str(audit.get("source_relative_path") or source.name)
        if not source.exists():
            unresolved_items.append({"source_relative_path": source_rel, "block_id": "", "quote": "", "problem": "source file missing", "recommendation": "", "status_note": str(source)})
            continue
        if source.suffix.lower() != ".docx":
            continue

        candidates: list[dict[str, Any]] = []
        local_unresolved: list[dict[str, Any]] = []
        for issue in audit.get("issues", []) or []:
            if not isinstance(issue, dict):
                continue
            candidate, unresolved = issue_to_candidate(issue, source_rel)
            if candidate is None:
                local_unresolved.append(unresolved)
            else:
                candidates.append(candidate)

        out_path = normalized_output_path(output_dir, input_dir, source)
        result = apply_candidates(source, out_path, candidates, dry_run=dry_run)
        applied_items.extend(item for item in candidates if item.get("status") == "applied")
        unresolved_items.extend(local_unresolved)
        unresolved_items.extend(item for item in candidates if item.get("status") != "applied")
        document_results.append({"source_relative_path": source_rel, "audit": str(audit_path), **result})

    stamp = time.strftime("%Y-%m-%d_%H-%M-%S")
    payload = {
        "schema_version": 1,
        "created_at": stamp,
        "provider": provider or "all",
        "documents": len(document_results),
        "document_results": document_results,
        "applied_items": applied_items,
        "unresolved_items": unresolved_items,
        "dry_run": dry_run,
    }
    report_dir.mkdir(parents=True, exist_ok=True)
    patch_dir.mkdir(parents=True, exist_ok=True)
    write_json(report_dir / f"{stamp}__normalization_{provider or 'all'}.json", payload)
    write_report_xlsx(report_dir / f"{stamp}__normalization_{provider or 'all'}.xlsx", payload)
    write_report_docx(report_dir / f"{stamp}__normalization_{provider or 'all'}.docx", payload)
    write_report_md(report_dir / f"{stamp}__normalization_{provider or 'all'}.md", payload)
    write_json(patch_dir / f"{stamp}__normalization_patch_plan_{provider or 'all'}.json", payload)
    write_json(patch_dir / f"latest_normalization_patch_plan_{provider or 'all'}.json", payload)
    return payload


def main() -> int:
    parser = argparse.ArgumentParser(description="Normalize DOCX documents from existing audit logs.")
    parser.add_argument("--provider", choices=["openai", "gemini", "xai", "anthropic", "all"], default="all")
    parser.add_argument("--from-logs", default="logs")
    parser.add_argument("--input", default="input")
    parser.add_argument("--output", default="output")
    parser.add_argument("--report-dir", default="output/_normalization")
    parser.add_argument("--patch-dir", default="report/document_normalization")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    root = THIS_DIR.parent
    normalize_from_audit_logs(
        root=root,
        provider="" if args.provider == "all" else args.provider,
        logs_dir=(root / args.from_logs).resolve() if not Path(args.from_logs).is_absolute() else Path(args.from_logs),
        input_dir=(root / args.input).resolve() if not Path(args.input).is_absolute() else Path(args.input),
        output_dir=(root / args.output).resolve() if not Path(args.output).is_absolute() else Path(args.output),
        report_dir=(root / args.report_dir).resolve() if not Path(args.report_dir).is_absolute() else Path(args.report_dir),
        patch_dir=(root / args.patch_dir).resolve() if not Path(args.patch_dir).is_absolute() else Path(args.patch_dir),
        dry_run=bool(args.dry_run),
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
