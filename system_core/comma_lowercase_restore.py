from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any
import argparse
import hashlib
import json
import re

from docx import Document


INVALID_REPORT_NAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
SPACE_RE = re.compile(r"\s+")
LLM_SYSTEM_PROMPT = """Ты проверяешь ложные правки после автоматического понижения регистра после запятой в DOCX.
Для каждого срабатывания реши, надо ли вернуть исходную заглавную букву.

Верни только JSON-объект:
{
  "decisions": [
    {"file": "имя.docx", "hit": 1, "action": "restore", "reason": "топоним/название/бренд"},
    {"file": "имя.docx", "hit": 2, "action": "keep", "reason": "обычное слово в перечислении"}
  ]
}

Правила:
- restore: топонимы, населённые пункты, районы как часть официального названия, бренды, названия объектов, собственные имена.
- keep: обычные нарицательные слова после запятой, элементы обычного перечисления, где заглавная буква была ошибочной.
- Не придумывай новые слова и не переписывай фрагменты текста.
- Если сомневаешься, выбирай keep и укажи причину сомнения.
"""


@dataclass(frozen=True)
class RestoreEntry:
    find: str
    replace: str


@dataclass(frozen=True)
class HitDecision:
    action: str
    replacement: str
    reason: str


def read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8-sig", errors="replace"))
    if not isinstance(payload, dict):
        raise RuntimeError(f"JSON report must contain an object: {path}")
    return payload


def one_line(text: str) -> str:
    return SPACE_RE.sub(" ", str(text or "")).strip()


def normalize_key(text: str) -> str:
    return one_line(text).casefold()


def parse_pair(text: str) -> tuple[str, str] | None:
    line = str(text).strip()
    if not line or line.startswith("#"):
        return None
    for separator in ("->", "=>", "=", "\t", ";"):
        if separator in line:
            left, right = line.split(separator, 1)
            left = left.strip()
            right = right.strip()
            return (left, right) if left and right else None
    if ":" in line:
        left, right = line.split(":", 1)
        left = left.strip()
        right = right.strip()
        return (left, right) if left and right else None
    return (line.casefold(), line)


def add_entry(entries: dict[str, RestoreEntry], find: Any, replace: Any) -> None:
    find_text = str(find or "").strip()
    replace_text = str(replace or "").strip()
    if not find_text or not replace_text:
        return
    entries[normalize_key(find_text)] = RestoreEntry(find=find_text, replace=replace_text)


def add_payload_entries(entries: dict[str, RestoreEntry], payload: Any) -> None:
    if payload is None:
        return
    if isinstance(payload, dict):
        known_sections = {
            "restore_words",
            "words",
            "restore_phrases",
            "phrases",
            "restore_map",
            "map",
            "replacements",
        }
        if any(key in payload for key in known_sections):
            for key in known_sections:
                add_payload_entries(entries, payload.get(key))
            return
        for find, replace in payload.items():
            add_entry(entries, find, replace)
        return
    if isinstance(payload, list):
        for item in payload:
            if isinstance(item, dict):
                if "from" in item or "to" in item:
                    add_entry(entries, item.get("from"), item.get("to"))
                elif "find" in item or "replace" in item:
                    add_entry(entries, item.get("find"), item.get("replace"))
                elif len(item) == 1:
                    find, replace = next(iter(item.items()))
                    add_entry(entries, find, replace)
                else:
                    add_payload_entries(entries, item)
            else:
                pair = parse_pair(str(item))
                if pair:
                    add_entry(entries, pair[0], pair[1])
        return
    pair = parse_pair(str(payload))
    if pair:
        add_entry(entries, pair[0], pair[1])


def load_restore_map(path: Path | None) -> list[RestoreEntry]:
    if path is None or not str(path).strip() or not path.exists():
        return []

    text = path.read_text(encoding="utf-8-sig", errors="replace")
    entries: dict[str, RestoreEntry] = {}
    suffix = path.suffix.lower()
    if suffix == ".json":
        add_payload_entries(entries, json.loads(text))
    elif suffix in {".yaml", ".yml"}:
        import yaml

        add_payload_entries(entries, yaml.safe_load(text))
    else:
        for line in text.splitlines():
            pair = parse_pair(line)
            if pair:
                add_entry(entries, pair[0], pair[1])
    return sorted(entries.values(), key=lambda entry: len(entry.find), reverse=True)


def decision_key(file_value: Any, hit_value: Any) -> tuple[str, int]:
    try:
        hit_id = int(hit_value or 0)
    except (TypeError, ValueError):
        hit_id = 0
    return (Path(str(file_value or "")).name.casefold(), hit_id)


def load_decisions(path: Path | None) -> dict[tuple[str, int], HitDecision]:
    if path is None or not str(path).strip() or not path.exists():
        return {}
    payload = read_json(path)
    raw_decisions = payload.get("decisions") or payload.get("items") or payload.get("hits") or []
    if isinstance(raw_decisions, dict):
        raw_decisions = [
            {"hit": key, **value} if isinstance(value, dict) else {"hit": key, "action": value}
            for key, value in raw_decisions.items()
        ]
    if not isinstance(raw_decisions, list):
        raise RuntimeError(f"Decisions JSON must contain a decisions/items list: {path}")

    decisions: dict[tuple[str, int], HitDecision] = {}
    default_file = payload.get("file") or payload.get("source") or payload.get("document") or ""
    for item in raw_decisions:
        if not isinstance(item, dict):
            continue
        action = str(item.get("action") or item.get("decision") or "").strip().lower()
        if not action and isinstance(item.get("restore"), bool):
            action = "restore" if item.get("restore") else "keep"
        if action in {"fix", "return", "revert", "исправить", "восстановить"}:
            action = "restore"
        if action in {"leave", "skip", "оставить", "не менять"}:
            action = "keep"
        if action not in {"restore", "keep"}:
            continue
        file_value = item.get("file") or item.get("document") or item.get("source") or default_file
        hit_value = item.get("hit") or item.get("hit_id") or item.get("index")
        key = decision_key(file_value, hit_value)
        if key[1] <= 0:
            continue
        decisions[key] = HitDecision(
            action=action,
            replacement=str(item.get("replacement") or item.get("restore_to") or "").strip(),
            reason=str(item.get("reason") or item.get("comment") or "llm_decision").strip(),
        )
    return decisions


def merge_decisions(
    base: dict[tuple[str, int], HitDecision],
    extra: dict[tuple[str, int], HitDecision],
) -> dict[tuple[str, int], HitDecision]:
    merged = dict(base)
    merged.update(extra)
    return merged


def paragraph_text_and_map(paragraph: Any) -> tuple[str, list[tuple[int, int]]]:
    chunks: list[str] = []
    char_map: list[tuple[int, int]] = []
    for run_index, run in enumerate(paragraph.runs):
        text = run.text or ""
        chunks.append(text)
        char_map.extend((run_index, char_index) for char_index in range(len(text)))
    return "".join(chunks), char_map


def table_by_path(document: Any, table_path: str) -> Any:
    parts = [int(part) for part in str(table_path or "").split(".") if part]
    if not parts:
        raise IndexError("empty table path")
    table = document.tables[parts[0] - 1]
    index = 1
    while index < len(parts):
        row_index, cell_index, nested_index = parts[index : index + 3]
        table = table.rows[row_index - 1].cells[cell_index - 1].tables[nested_index - 1]
        index += 3
    return table


def paragraph_for_hit(document: Any, hit: dict[str, Any]) -> Any:
    kind = str(hit.get("kind") or "")
    paragraph_index = int(hit.get("paragraph") or 0)
    if paragraph_index <= 0:
        raise IndexError("invalid paragraph index")
    if kind == "paragraph":
        return document.paragraphs[paragraph_index - 1]
    if kind == "table_cell":
        table = table_by_path(document, str(hit.get("table") or ""))
        row = table.rows[int(hit.get("row") or 0) - 1]
        cell = row.cells[int(hit.get("cell") or 0) - 1]
        return cell.paragraphs[paragraph_index - 1]
    raise IndexError(f"unsupported hit kind: {kind}")


def same_text_ignoring_case(left: str, right: str) -> bool:
    return str(left).casefold() == str(right).casefold()


def replace_same_length(paragraph: Any, start: int, current: str, replacement: str) -> bool:
    if len(current) != len(replacement):
        return False
    text, char_map = paragraph_text_and_map(paragraph)
    end = start + len(current)
    if start < 0 or end > len(text) or end > len(char_map):
        return False
    if text[start:end] != current:
        return False

    run_chars = [list(run.text or "") for run in paragraph.runs]
    for offset, char in enumerate(replacement):
        run_index, char_index = char_map[start + offset]
        run_chars[run_index][char_index] = char
    for run, chars in zip(paragraph.runs, run_chars):
        run.text = "".join(chars)
    return True


def iter_hits(result: dict[str, Any]) -> list[dict[str, Any]]:
    hits = result.get("hits")
    if isinstance(hits, list):
        return [hit for hit in hits if isinstance(hit, dict)]

    out: list[dict[str, Any]] = []
    for location in result.get("locations", []) or []:
        if isinstance(location, dict):
            out.extend(hit for hit in location.get("hits", []) or [] if isinstance(hit, dict))
    return out


def format_location(hit: dict[str, Any]) -> str:
    if hit.get("kind") == "table_cell":
        return (
            f"таблица {hit.get('table')}, строка {hit.get('row')}, "
            f"ячейка {hit.get('cell')}, абзац {hit.get('paragraph')}"
        )
    if hit.get("kind") == "paragraph":
        return f"абзац {hit.get('paragraph')}"
    return str(hit.get("kind") or "document")


def build_llm_cards(results: list[dict[str, Any]], scope: str) -> list[dict[str, Any]]:
    cards: list[dict[str, Any]] = []
    for result in results:
        file_name = Path(str(result.get("file") or "document.docx")).name
        for hit in iter_hits(result):
            if scope == "table-cells" and hit.get("kind") != "table_cell":
                continue
            cards.append(
                {
                    "file": file_name,
                    "hit": int(hit.get("hit") or 0),
                    "where": format_location(hit),
                    "original_word": str(hit.get("word_before") or ""),
                    "changed_word": str(hit.get("word_after") or ""),
                    "before_context": one_line(str(hit.get("before") or "")),
                    "after_context": one_line(str(hit.get("after") or "")),
                    "after_fragment": one_line(str(hit.get("after_fragment") or "")),
                    "following_text": one_line(str(hit.get("following_text") or "")),
                }
            )
    return [card for card in cards if card["hit"] > 0 and card["changed_word"]]


def chunked(items: list[dict[str, Any]], size: int) -> list[list[dict[str, Any]]]:
    size = max(1, int(size or 1))
    return [items[index : index + size] for index in range(0, len(items), size)]


def decisions_from_payload(payload: dict[str, Any]) -> dict[tuple[str, int], HitDecision]:
    temp_path = None
    decisions: dict[tuple[str, int], HitDecision] = {}
    raw_decisions = payload.get("decisions") if isinstance(payload, dict) else []
    if not isinstance(raw_decisions, list):
        return decisions
    for item in raw_decisions:
        if not isinstance(item, dict):
            continue
        action = str(item.get("action") or item.get("decision") or "").strip().lower()
        if action in {"fix", "return", "revert", "исправить", "восстановить"}:
            action = "restore"
        if action in {"leave", "skip", "оставить", "не менять"}:
            action = "keep"
        if action not in {"restore", "keep"}:
            continue
        key = decision_key(item.get("file") or temp_path or "", item.get("hit") or item.get("hit_id") or item.get("index"))
        if key[1] <= 0:
            continue
        decisions[key] = HitDecision(
            action=action,
            replacement=str(item.get("replacement") or item.get("restore_to") or "").strip(),
            reason=str(item.get("reason") or item.get("comment") or "llm_decision").strip(),
        )
    return decisions


def call_llm_for_decisions(
    results: list[dict[str, Any]],
    *,
    provider: str,
    model: str,
    key_ref: str,
    scope: str,
    batch_size: int,
    max_output_tokens: int,
    max_retries: int,
    timeout_sec: float,
    reasoning_effort: str,
    service_tier: str,
) -> tuple[dict[tuple[str, int], HitDecision], dict[str, Any]]:
    provider = provider.strip().lower()
    cards = build_llm_cards(results, scope)
    batches = chunked(cards, batch_size)
    usage_total = {"calls": 0, "input_tokens": 0, "output_tokens": 0, "total_tokens": 0, "reasoning_tokens": 0}
    decisions: dict[tuple[str, int], HitDecision] = {}
    raw_batches: list[dict[str, Any]] = []

    if not cards:
        return decisions, {"cards": 0, "batches": 0, "usage": usage_total, "raw_batches": raw_batches}

    from config_resolver import load_settings, resolve_api_key, resolve_model

    settings = load_settings()
    if not model:
        if provider == "openai":
            model = resolve_model("openai", "audit", settings)
        elif provider == "gemini":
            model = resolve_model("gemini", "audit_fast", settings)
        elif provider == "xai":
            model = resolve_model("xai", "audit", settings)
        elif provider == "anthropic":
            model = resolve_model("anthropic", "audit", settings)
    if not model:
        raise RuntimeError(f"No model configured for LLM decisions provider: {provider}")

    if provider == "openai":
        from openai import OpenAI
        from providers.openai_provider import call_json_object

        api_key = resolve_api_key("openai", settings, key_ref=key_ref)
        if not api_key:
            raise RuntimeError("OpenAI API key not found. Set OPENAI_API_KEY or config\\api_key_openai.txt.")
        client = OpenAI(api_key=api_key, timeout=timeout_sec, max_retries=0)
        for index, batch in enumerate(batches, start=1):
            user_prompt = json.dumps({"hits": batch}, ensure_ascii=False, indent=2)
            doc_hash = hashlib.sha256(user_prompt.encode("utf-8", errors="replace")).hexdigest()
            print(f"[LLM DECISIONS] OpenAI batch {index}/{len(batches)} hits={len(batch)}")
            obj, usage, tier = call_json_object(
                client,
                model=model,
                instructions=LLM_SYSTEM_PROMPT,
                user_prompt=user_prompt,
                reasoning_effort=reasoning_effort,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max(1, int(max_retries)),
                service_tier=service_tier,
                use_idempotency=True,
                doc_hash=doc_hash,
                chunk_index=index,
            )
            decisions = merge_decisions(decisions, decisions_from_payload(obj))
            raw_batches.append({"index": index, "provider": provider, "model": model, "response": obj, "usage": usage, "service_tier": tier})
            usage_total["calls"] += 1
            for key in ("input_tokens", "output_tokens", "total_tokens", "reasoning_tokens"):
                usage_total[key] += int(usage.get(key, 0) or 0)
    elif provider == "gemini":
        from google import genai
        from providers.gemini_provider import call_structured

        api_key = resolve_api_key("gemini", settings, key_ref=key_ref)
        if not api_key:
            raise RuntimeError("Gemini API key not found. Set GEMINI_API_KEY or config\\api_key_gemini.txt.")
        client = genai.Client(api_key=api_key)
        for index, batch in enumerate(batches, start=1):
            user_prompt = json.dumps({"hits": batch}, ensure_ascii=False, indent=2)
            print(f"[LLM DECISIONS] Gemini batch {index}/{len(batches)} hits={len(batch)}")
            obj, usage, tier = call_structured(
                client,
                model=model,
                system_instruction=LLM_SYSTEM_PROMPT,
                user_prompt=user_prompt,
                max_retries=max(1, int(max_retries)),
            )
            decisions = merge_decisions(decisions, decisions_from_payload(obj))
            raw_batches.append({"index": index, "provider": provider, "model": model, "response": obj, "usage": usage, "service_tier": tier})
            usage_total["calls"] += 1
            for key in ("input_tokens", "output_tokens", "total_tokens", "reasoning_tokens"):
                usage_total[key] += int(usage.get(key, 0) or 0)
    elif provider == "xai":
        from providers.xai_provider import call_json_object

        api_key = resolve_api_key("xai", settings, key_ref=key_ref)
        if not api_key:
            raise RuntimeError("xAI API key not found. Set XAI_API_KEY or config\\api_key_xai.txt.")
        for index, batch in enumerate(batches, start=1):
            user_prompt = json.dumps({"hits": batch}, ensure_ascii=False, indent=2)
            doc_hash = hashlib.sha256(user_prompt.encode("utf-8", errors="replace")).hexdigest()
            print(f"[LLM DECISIONS] xAI batch {index}/{len(batches)} hits={len(batch)}")
            obj, usage, tier = call_json_object(
                api_key=api_key,
                model=model,
                instructions=LLM_SYSTEM_PROMPT,
                user_prompt=user_prompt,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max(1, int(max_retries)),
                use_idempotency=True,
                doc_hash=doc_hash,
                chunk_index=index,
            )
            decisions = merge_decisions(decisions, decisions_from_payload(obj))
            raw_batches.append({"index": index, "provider": provider, "model": model, "response": obj, "usage": usage, "service_tier": tier})
            usage_total["calls"] += 1
            for key in ("input_tokens", "output_tokens", "total_tokens", "reasoning_tokens"):
                usage_total[key] += int(usage.get(key, 0) or 0)
    elif provider == "anthropic":
        from providers.anthropic_provider import call_json_object

        api_key = resolve_api_key("anthropic", settings, key_ref=key_ref)
        if not api_key:
            raise RuntimeError("Anthropic API key not found. Set ANTHROPIC_API_KEY or config\\api_key_anthropic.txt.")
        for index, batch in enumerate(batches, start=1):
            user_prompt = json.dumps({"hits": batch}, ensure_ascii=False, indent=2)
            print(f"[LLM DECISIONS] Claude batch {index}/{len(batches)} hits={len(batch)}")
            obj, usage, tier = call_json_object(
                api_key=api_key,
                model=model,
                instructions=LLM_SYSTEM_PROMPT,
                user_prompt=user_prompt,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max(1, int(max_retries)),
                reasoning_effort=reasoning_effort,
            )
            decisions = merge_decisions(decisions, decisions_from_payload(obj))
            raw_batches.append({"index": index, "provider": provider, "model": model, "response": obj, "usage": usage, "service_tier": tier})
            usage_total["calls"] += 1
            for key in ("input_tokens", "output_tokens", "total_tokens", "reasoning_tokens"):
                usage_total[key] += int(usage.get(key, 0) or 0)
    else:
        raise RuntimeError(f"Unsupported LLM decisions provider: {provider}")

    return decisions, {"cards": len(cards), "batches": len(batches), "usage": usage_total, "raw_batches": raw_batches}


def report_relative_path(payload: dict[str, Any], result: dict[str, Any]) -> Path:
    source = Path(str(result.get("file") or "document.docx"))
    input_root = Path(str(payload.get("input") or ""))
    try:
        return source.relative_to(input_root)
    except (ValueError, OSError):
        return Path(source.name)


def unique_by_name(root: Path, name: str) -> Path | None:
    if not root.is_dir():
        return None
    matches = [path for path in root.rglob(name) if path.is_file()]
    return matches[0] if len(matches) == 1 else None


def find_source_doc(payload: dict[str, Any], result: dict[str, Any], source: Path) -> Path | None:
    if source.is_file():
        result_name = Path(str(result.get("file") or "")).name
        return source if not result_name or source.name == result_name else None
    relative = report_relative_path(payload, result)
    candidates = [
        source / relative,
        source / relative.name,
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return unique_by_name(source, relative.name)


def find_fixed_doc(payload: dict[str, Any], result: dict[str, Any], fixed: Path) -> Path | None:
    if fixed.is_file():
        return fixed
    relative = report_relative_path(payload, result)
    source_name = relative.name
    source_stem = Path(source_name).stem
    output_value = str(result.get("output") or "")
    output_root = Path(str(payload.get("output") or ""))
    candidates: list[Path] = []
    if output_value:
        output_path = Path(output_value)
        try:
            candidates.append(fixed / output_path.relative_to(output_root))
        except (ValueError, OSError):
            candidates.append(fixed / output_path.name)
    candidates.extend(
        [
            fixed / relative.parent / f"{source_stem}_comma_lowercase.docx",
            fixed / relative,
            fixed / source_name,
            fixed / f"{source_stem}_comma_lowercase.docx",
        ]
    )
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return unique_by_name(fixed, f"{source_stem}_comma_lowercase.docx") or unique_by_name(fixed, source_name)


def unique_path(path: Path, overwrite: bool) -> Path:
    if overwrite or not path.exists():
        return path
    for index in range(2, 1000):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Could not allocate output path: {path}")


def output_path_for(source_doc: Path, fixed_doc: Path, fixed_root: Path, output: Path, multiple: bool, overwrite: bool) -> Path:
    if output.suffix.lower() == ".docx" and not multiple:
        return unique_path(output, overwrite)
    try:
        relative = fixed_doc.relative_to(fixed_root if fixed_root.is_dir() else fixed_doc.parent)
    except ValueError:
        relative = Path(fixed_doc.name)
    stem = relative.stem
    if stem == source_doc.stem:
        stem = f"{stem}_comma_restored"
    elif not stem.endswith("_restored"):
        stem = f"{stem}_restored"
    return unique_path(output / relative.parent / f"{stem}.docx", overwrite)


def safe_report_stem(path: Path) -> str:
    stem = INVALID_REPORT_NAME_RE.sub("_", path.stem).strip(" .")
    return stem or "document"


def select_result_for_single_source(payload: dict[str, Any], source: Path) -> dict[str, Any] | None:
    results = [result for result in payload.get("results", []) or [] if isinstance(result, dict)]
    if not results:
        return None
    by_name = [result for result in results if Path(str(result.get("file") or "")).name == source.name]
    if len(by_name) == 1:
        return by_name[0]
    changed = [result for result in results if int(result.get("changes") or 0) > 0]
    return changed[0] if len(changed) == 1 else None


def choose_restore(
    hit: dict[str, Any],
    fixed_text: str,
    position: int,
    restore_entries: list[RestoreEntry],
    decision: HitDecision | None,
    restore_all: bool,
) -> tuple[str, str, str] | None:
    if decision is not None:
        if decision.action == "keep":
            return ("", "", "llm_keep")
        current = str(hit.get("word_after") or "")
        replacement = decision.replacement or str(hit.get("word_before") or "")
        if current and replacement and fixed_text[position : position + len(current)] == current:
            return current, replacement, f"llm_restore: {decision.reason}"
        return None

    if restore_all:
        current = str(hit.get("word_after") or "")
        replacement = str(hit.get("word_before") or "")
        if current and replacement and fixed_text[position : position + len(current)] == current:
            return current, replacement, "restore_all"
        return None

    for entry in restore_entries:
        current = fixed_text[position : position + len(entry.find)]
        if len(current) != len(entry.find):
            continue
        if normalize_key(current) == normalize_key(entry.find):
            return current, entry.replace, "restore_map"
    return None


def process_document(
    payload: dict[str, Any],
    result: dict[str, Any],
    source_doc: Path,
    fixed_doc: Path,
    output_doc: Path,
    restore_entries: list[RestoreEntry],
    decisions: dict[tuple[str, int], HitDecision],
    *,
    scope: str,
    restore_all: bool,
    dry_run: bool,
) -> dict[str, Any]:
    fixed_document = Document(str(fixed_doc))
    source_document = Document(str(source_doc)) if source_doc.is_file() else None
    restored = 0
    skipped = 0
    items: list[dict[str, Any]] = []

    for hit in iter_hits(result):
        location = {
            key: hit.get(key)
            for key in ("kind", "table", "row", "cell", "paragraph", "hit", "position", "word_before", "word_after")
            if key in hit
        }
        if scope == "table-cells" and hit.get("kind") != "table_cell":
            skipped += 1
            items.append({"status": "skipped", "reason": "outside_table_cells", "location": location})
            continue

        try:
            position = int(hit.get("position") or 0)
            fixed_paragraph = paragraph_for_hit(fixed_document, hit)
            fixed_text, _ = paragraph_text_and_map(fixed_paragraph)
        except Exception as exc:
            skipped += 1
            items.append({"status": "skipped", "reason": f"fixed_location_error: {exc}", "location": location})
            continue

        llm_decision = decisions.get(decision_key(result.get("file"), hit.get("hit"))) or decisions.get(decision_key("", hit.get("hit")))
        decision = choose_restore(hit, fixed_text, position, restore_entries, llm_decision, restore_all)
        if decision is not None and decision[2] == "llm_keep":
            skipped += 1
            items.append({"status": "skipped", "reason": "llm_keep", "location": location})
            continue
        if decision is None:
            skipped += 1
            reason = "no_restore_decision" if decisions else "not_in_restore_map"
            items.append({"status": "skipped", "reason": reason, "location": location})
            continue
        current, replacement, source_mode = decision

        source_segment = ""
        if source_document is not None:
            try:
                source_paragraph = paragraph_for_hit(source_document, hit)
                source_text, _ = paragraph_text_and_map(source_paragraph)
                source_segment = source_text[position : position + len(current)]
            except Exception as exc:
                skipped += 1
                items.append({"status": "skipped", "reason": f"source_location_error: {exc}", "location": location})
                continue
            if not same_text_ignoring_case(source_segment, current):
                skipped += 1
                items.append(
                    {
                        "status": "skipped",
                        "reason": "source_fixed_text_mismatch",
                        "location": location,
                        "current": current,
                        "source": source_segment,
                    }
                )
                continue
            replacement = source_segment

        if len(current) != len(replacement):
            skipped += 1
            items.append(
                {
                    "status": "skipped",
                    "reason": "replacement_length_mismatch",
                    "location": location,
                    "current": current,
                    "replacement": replacement,
                }
            )
            continue

        if current == replacement:
            skipped += 1
            items.append(
                {
                    "status": "skipped",
                    "reason": "already_restored",
                    "location": location,
                    "current": current,
                    "replacement": replacement,
                }
            )
            continue

        if not dry_run and not replace_same_length(fixed_paragraph, position, current, replacement):
            skipped += 1
            items.append(
                {
                    "status": "skipped",
                    "reason": "fixed_text_changed_before_apply",
                    "location": location,
                    "current": current,
                    "replacement": replacement,
                }
            )
            continue

        restored += 1
        items.append(
            {
                "status": "restored" if not dry_run else "would_restore",
                "reason": source_mode,
                "location": location,
                "current": current,
                "replacement": replacement,
                "source": source_segment,
            }
        )

    if restored and not dry_run:
        output_doc.parent.mkdir(parents=True, exist_ok=True)
        fixed_document.save(str(output_doc))

    return {
        "source": str(source_doc),
        "fixed": str(fixed_doc),
        "output": str(output_doc) if restored and not dry_run else "",
        "report_file": str(result.get("file") or ""),
        "hits": len(iter_hits(result)),
        "restored": restored,
        "skipped": skipped,
        "items": items,
    }


def render_markdown_report(report: dict[str, Any]) -> str:
    lines = [
        "# Восстановление регистра после запятых",
        "",
        f"- Файлов: {report['files']}",
        f"- Восстановлено: {report['restored']}",
        f"- Пропущено: {report['skipped']}",
        f"- Dry-run: {'да' if report['dry_run'] else 'нет'}",
        f"- Область: {'только ячейки таблиц' if report['scope'] == 'table-cells' else 'все срабатывания отчёта'}",
        f"- Карта восстановления: {report['restore_map'] or 'не задана'}",
        "",
    ]
    for result in report["results"]:
        lines.extend(
            [
                f"## {Path(str(result['fixed'])).name}",
                "",
                f"- Срабатываний в отчёте: {result['hits']}",
                f"- Восстановлено: {result['restored']}",
                f"- Пропущено: {result['skipped']}",
            ]
        )
        if result.get("output"):
            lines.append(f"- Документ: {result['output']}")
        lines.append("")

        restored_items = [item for item in result["items"] if item.get("status") in {"restored", "would_restore"}]
        if not restored_items:
            lines.extend(["Восстановлений нет.", ""])
            continue
        for item in restored_items:
            location = item.get("location", {})
            if location.get("kind") == "table_cell":
                where = (
                    f"таблица {location.get('table')}, строка {location.get('row')}, "
                    f"ячейка {location.get('cell')}, абзац {location.get('paragraph')}"
                )
            else:
                where = f"абзац {location.get('paragraph')}"
            lines.append(f"- {where}: `{item.get('current')}` -> `{item.get('replacement')}`")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def build_pairs(payload: dict[str, Any], source: Path, fixed: Path, output: Path, overwrite: bool) -> list[tuple[dict[str, Any], Path, Path, Path]]:
    results = [result for result in payload.get("results", []) or [] if isinstance(result, dict)]
    pairs: list[tuple[dict[str, Any], Path, Path, Path]] = []
    if source.is_file() and fixed.is_file():
        result = select_result_for_single_source(payload, source)
        if result is None:
            raise RuntimeError("Could not choose a report result for the provided source DOCX.")
        pairs.append((result, source, fixed, output_path_for(source, fixed, fixed, output, False, overwrite)))
        return pairs

    for result in results:
        if not iter_hits(result):
            continue
        source_doc = find_source_doc(payload, result, source)
        fixed_doc = find_fixed_doc(payload, result, fixed)
        if source_doc is None or fixed_doc is None:
            continue
        out_doc = output_path_for(source_doc, fixed_doc, fixed, output, True, overwrite)
        pairs.append((result, source_doc, fixed_doc, out_doc))
    return pairs


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Restore selected false comma-lowercase DOCX edits using a DocFlow JSON report and a restore map.",
    )
    parser.add_argument("--report", required=True, help="DocFlow comma_lowercase JSON report.")
    parser.add_argument("--source", required=True, help="Original DOCX file or folder.")
    parser.add_argument("--fixed", required=True, help="Corrected comma-lowercase DOCX file or folder.")
    parser.add_argument("--output", default="output/comma_lowercase_restored", help="Output DOCX file or folder.")
    parser.add_argument("--restore-map", default="config/comma_restore_map.yaml", help="YAML/JSON/TXT restore map.")
    parser.add_argument("--decisions", default="", help="Optional LLM decisions JSON with restore/keep per report hit.")
    parser.add_argument("--llm-provider", choices=("none", "openai", "gemini", "xai", "anthropic"), default="none", help="Optionally ask an LLM to classify report hits.")
    parser.add_argument("--llm-model", default="", help="Manual model override for LLM decisions.")
    parser.add_argument("--openai-api-key-ref", default="", help="OpenAI API key ref from config/api_key_openai.txt.")
    parser.add_argument("--gemini-api-key-ref", default="", help="Gemini API key ref from config/api_key_gemini.txt.")
    parser.add_argument("--xai-api-key-ref", default="", help="xAI API key ref from config/api_key_xai.txt.")
    parser.add_argument("--anthropic-api-key-ref", default="", help="Anthropic API key ref from config/api_key_anthropic.txt.")
    parser.add_argument("--llm-batch-size", type=int, default=80, help="Hits per LLM decision batch.")
    parser.add_argument("--llm-max-output-tokens", type=int, default=8000, help="Max LLM output tokens per batch.")
    parser.add_argument("--llm-max-retries", type=int, default=2, help="LLM retries per batch.")
    parser.add_argument("--llm-timeout-sec", type=float, default=300.0, help="OpenAI timeout seconds.")
    parser.add_argument("--llm-reasoning-effort", default="low", help="OpenAI reasoning effort for decisions.")
    parser.add_argument("--llm-service-tier", default="auto", help="OpenAI service tier.")
    parser.add_argument("--decisions-out", default="", help="Where to write generated/merged LLM decisions JSON.")
    parser.add_argument("--json-out", default="report/comma_lowercase_restore.json", help="JSON restore report path.")
    parser.add_argument("--md-out", default="", help="Markdown restore report path.")
    parser.add_argument("--scope", choices=("table-cells", "all"), default="table-cells", help="Which report hits may be restored.")
    parser.add_argument("--restore-all", action="store_true", help="Restore every selected report hit. For diagnostics only.")
    parser.add_argument("--dry-run", action="store_true", help="Report planned restorations without writing DOCX files.")
    parser.add_argument("--overwrite", action="store_true", help="Allow overwriting explicit output files.")
    return parser


def run(args: argparse.Namespace) -> int:
    report_path = Path(args.report).resolve()
    source = Path(args.source).resolve()
    fixed = Path(args.fixed).resolve()
    output = Path(args.output).resolve()
    restore_map = Path(args.restore_map).resolve() if str(args.restore_map or "").strip() else None
    decisions_path = Path(args.decisions).resolve() if str(args.decisions or "").strip() else None
    json_out = Path(args.json_out).resolve()
    md_out = Path(args.md_out).resolve() if str(args.md_out or "").strip() else json_out.with_suffix(".md")
    decisions_out = Path(args.decisions_out).resolve() if str(args.decisions_out or "").strip() else json_out.with_name(f"{json_out.stem}_decisions.json")

    payload = read_json(report_path)
    restore_entries = load_restore_map(restore_map)
    decisions = load_decisions(decisions_path)

    pairs = build_pairs(payload, source, fixed, output, bool(args.overwrite))
    if not pairs:
        raise RuntimeError("No matching source/fixed DOCX pairs were found for the report.")

    llm_payload: dict[str, Any] = {}
    provider = str(args.llm_provider or "none").strip().lower()
    if provider != "none":
        llm_decisions, llm_payload = call_llm_for_decisions(
            [pair[0] for pair in pairs],
            provider=provider,
            model=str(args.llm_model or "").strip(),
            key_ref=str(
                {
                    "openai": args.openai_api_key_ref,
                    "gemini": args.gemini_api_key_ref,
                    "xai": args.xai_api_key_ref,
                    "anthropic": args.anthropic_api_key_ref,
                }.get(provider, "")
                or ""
            ).strip(),
            scope=str(args.scope),
            batch_size=int(args.llm_batch_size),
            max_output_tokens=int(args.llm_max_output_tokens),
            max_retries=int(args.llm_max_retries),
            timeout_sec=float(args.llm_timeout_sec),
            reasoning_effort=str(args.llm_reasoning_effort or "low"),
            service_tier=str(args.llm_service_tier or "auto"),
        )
        decisions = merge_decisions(decisions, llm_decisions)
        decisions_out.parent.mkdir(parents=True, exist_ok=True)
        decisions_out.write_text(
            json.dumps(
                {
                    "provider": provider,
                    "model": str(args.llm_model or ""),
                    "source_report": str(report_path),
                    "decisions": [
                        {
                            "file": file_name,
                            "hit": hit_id,
                            "action": decision.action,
                            "replacement": decision.replacement,
                            "reason": decision.reason,
                        }
                        for (file_name, hit_id), decision in sorted(decisions.items())
                    ],
                    "llm": llm_payload,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

    if not restore_entries and not decisions and not args.restore_all:
        print("Restore map and decisions are empty; no hits will be restored.")

    results = [
        process_document(
            payload,
            result,
            source_doc,
            fixed_doc,
            output_doc,
            restore_entries,
            decisions,
            scope=str(args.scope),
            restore_all=bool(args.restore_all),
            dry_run=bool(args.dry_run),
        )
        for result, source_doc, fixed_doc, output_doc in pairs
    ]
    report = {
        "command": "comma-lowercase-restore",
        "dry_run": bool(args.dry_run),
        "report": str(report_path),
        "source": str(source),
        "fixed": str(fixed),
        "output": str(output),
        "restore_map": str(restore_map) if restore_map else "",
        "decisions": str(decisions_path) if decisions_path else "",
        "decisions_out": str(decisions_out) if provider != "none" else "",
        "llm_provider": provider,
        "llm": llm_payload,
        "scope": str(args.scope),
        "restore_all": bool(args.restore_all),
        "files": len(results),
        "restored": sum(int(result["restored"]) for result in results),
        "skipped": sum(int(result["skipped"]) for result in results),
        "results": results,
    }
    json_out.parent.mkdir(parents=True, exist_ok=True)
    json_out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    md_out.parent.mkdir(parents=True, exist_ok=True)
    md_out.write_text(render_markdown_report(report), encoding="utf-8")

    print(f"Files: {report['files']}; restored: {report['restored']}; skipped: {report['skipped']}")
    print(f"JSON report: {json_out}")
    print(f"Markdown report: {md_out}")
    if args.dry_run:
        print("Dry-run: DOCX files were not written.")
    return 0


def main(argv: list[str] | None = None) -> int:
    return run(build_parser().parse_args(argv))


if __name__ == "__main__":
    raise SystemExit(main())
