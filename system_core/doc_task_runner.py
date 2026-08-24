#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""General LLM document task runner for report and exact DOCX replacements."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sys
import threading
import time
from pathlib import Path
from typing import Any, Dict, List

THIS_DIR = Path(__file__).resolve().parent
if str(THIS_DIR) not in sys.path:
    sys.path.insert(0, str(THIS_DIR))

from config_resolver import CONFIG_DIR, load_settings, resolve_api_key, resolve_model
from doc_task_resolver import TASK_QUERY_ENV, load_doc_task, remember_doc_task_use, remember_quick_doc_task_use
from doc_task_document_model import (
    block_lookup,
    block_paragraph_refs_docx,
    build_corpus_block_map,
    build_task_block_map,
    default_project_paths,
    ensure_project_dirs,
    iter_task_inputs,
    rel_parent,
    source_relative,
)
from chunk_workers import DEFAULT_WORKERS, resolve_workers, run_jobs
from llm_audit_helpers import build_chunks
from render.com_renderer import format_duration


SYSTEM_PROMPT_FILE = CONFIG_DIR / "doc_task_system_prompt.md"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig", errors="replace")


def write_json(path: Path, payload: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def read_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig", errors="replace"))


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(str(text or "").encode("utf-8", errors="ignore")).hexdigest()


def safe_name(text: str, max_len: int = 80) -> str:
    return "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in text)[:max_len] or "default"


def task_artifact_paths(
    paths,
    source: Path,
    provider: str,
    model: str,
    *,
    edited_suffix: str = "__doc_task",
) -> Dict[str, Path]:
    parent = rel_parent(source, paths.input_dir)
    stem = source.stem
    suffix = source.suffix.lower()
    safe_provider = safe_name(provider, 32)
    safe_model = safe_name(model, 80)
    return {
        "block_map": paths.logs_dir / parent / f"{stem}__doc_task_block_map.json",
        "task": paths.logs_dir / parent / f"{stem}__doc_task.json",
        "cache": paths.cache_dir / parent / f"{stem}__doc_task__{safe_provider}_{safe_model}.json",
        "edited": paths.output_dir / parent / f"{stem}{edited_suffix}{suffix}",
    }


def source_inventory_hash(sources: list[Path], input_dir: Path, *parts: str) -> str:
    h = hashlib.sha256()
    for part in parts:
        h.update(str(part or "").encode("utf-8", errors="ignore"))
        h.update(b"\0")
    for source in sources:
        rel = source_relative(source, input_dir).as_posix()
        try:
            stamp = f"{rel}|{sha256_file(source)}"
        except OSError:
            stamp = rel
        h.update(stamp.encode("utf-8", errors="ignore"))
        h.update(b"\0")
    return h.hexdigest()[:24]


def doc_task_cache_signature(
    *,
    source_sha256: str,
    task_ref: str,
    task_instruction: str,
    user_query: str,
    provider: str,
    model: str,
    chunk_tokens: int,
    overlap_tokens: int,
    min_chunks: int,
    max_output_tokens: int,
    pdf_max_pages: int,
    system_prompt_sha256: str,
) -> dict[str, Any]:
    return {
        "cache_version": 2,
        "source_sha256": str(source_sha256 or ""),
        "task_ref": str(task_ref or ""),
        "task_instruction_sha256": sha256_text(task_instruction),
        "user_query_sha256": sha256_text(user_query),
        "provider": str(provider or "").lower().strip(),
        "model": str(model or "").strip(),
        "chunk_tokens": int(chunk_tokens),
        "overlap_tokens": int(overlap_tokens),
        "min_chunks": int(min_chunks),
        "max_output_tokens": int(max_output_tokens),
        "pdf_max_pages": int(pdf_max_pages),
        "system_prompt_sha256": str(system_prompt_sha256 or ""),
    }


def doc_task_cache_mismatch_reason(payload: dict[str, Any], expected: dict[str, Any]) -> str:
    actual = payload.get("cache_signature")
    if not isinstance(actual, dict):
        return "missing cache_signature"
    for key, expected_value in expected.items():
        if actual.get(key) != expected_value:
            return f"{key} changed"
    return ""


def corpus_artifact_paths(paths, cache_id: str, provider: str, model: str) -> Dict[str, Path]:
    safe_provider = safe_name(provider, 32)
    safe_model = safe_name(model, 80)
    return {
        "block_map": paths.logs_dir / "_doc_tasks" / f"corpus__{cache_id}__block_map.json",
        "task": paths.logs_dir / "_doc_tasks" / f"corpus__{cache_id}__doc_task.json",
        "cache": paths.cache_dir / "_doc_tasks" / f"corpus__{cache_id}__{safe_provider}_{safe_model}.json",
    }


def document_item_from_block_map(block_map: Dict[str, Any]) -> dict[str, Any]:
    return {
        "source_relative_path": block_map.get("source_relative_path", ""),
        "document_type": block_map.get("document_type", ""),
        "blocks": len(block_map.get("blocks", []) or []),
    }


def resolve_task_scope(scope: str, source_count: int, apply_replacements: bool) -> str:
    normalized = str(scope or "auto").strip().lower()
    if normalized not in {"auto", "document", "corpus"}:
        normalized = "auto"
    if normalized == "auto":
        return "document" if apply_replacements or source_count <= 1 else "corpus"
    if normalized == "corpus" and apply_replacements:
        print("[DOC TASK] corpus scope cannot apply exact DOCX replacements safely; using document scope.")
        return "document"
    return normalized


def build_block_text(block: Dict[str, Any]) -> str:
    block_id = str(block.get("block_id") or "")
    kind = str(block.get("object_type") or block.get("kind") or "")
    source = str(block.get("source_relative_path") or "")
    location = str(block.get("location") or "")
    text = str(block.get("text") or "")
    meta = [f"[BLOCK:{block_id}]", f"type={kind}"]
    if source:
        meta.append(f"file={source}")
    if location:
        meta.append(f"location={location}")
    return f"{' '.join(meta)}\n{text}"


def build_corpus_context(block_map: Dict[str, Any]) -> str:
    documents = block_map.get("documents", []) or []
    if not documents:
        return "(single document)"
    lines = []
    for item in documents:
        if not isinstance(item, dict):
            continue
        lines.append(
            f"- d{int(item.get('document_index') or 0):03d}: "
            f"{item.get('source_relative_path', '')} "
            f"({item.get('document_type', '')}, blocks={item.get('blocks', 0)})"
        )
    return "\n".join(lines) or "(single document)"


def build_prompt(
    task_instruction: str,
    user_query: str,
    block_map: Dict[str, Any],
    overlap: list[dict[str, Any]],
    chunk: list[dict[str, Any]],
) -> str:
    overlap_text = "\n\n".join(build_block_text(block) for block in overlap)
    chunk_text = "\n\n".join(build_block_text(block) for block in chunk)
    return (
        "TASK INSTRUCTION:\n"
        f"{task_instruction.strip()}\n\n"
        "USER REQUEST:\n"
        f"{user_query.strip() or 'Follow the selected task instruction.'}\n\n"
        "OUTPUT FORMAT:\n"
        "Return exactly one JSON object matching the system prompt schema.\n\n"
        "CORPUS DOCUMENTS:\n"
        f"{build_corpus_context(block_map)}\n\n"
        "OVERLAP BLOCKS:\n"
        f"{overlap_text or '(none)'}\n\n"
        "NEW DOCUMENT BLOCKS:\n"
        f"{chunk_text}"
    )


def normalize_rows(
    rows: list[dict[str, Any]],
    source_rel: str,
    block_by_id: dict[str, dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for index, row in enumerate(rows, start=1):
        if not isinstance(row, dict):
            continue
        item = dict(row)
        item.setdefault("row_id", f"R{index:03d}")
        block = (block_by_id or {}).get(str(item.get("block_id") or ""))
        if not str(item.get("source_relative_path") or "").strip():
            item["source_relative_path"] = str((block or {}).get("source_relative_path") or source_rel)
        if not str(item.get("location") or "").strip() and block:
            item["location"] = str(block.get("location") or "")
        if not isinstance(item.get("values"), dict):
            item.pop("values", None)
        out.append(item)
    return out


def normalize_replacements(
    replacements: list[dict[str, Any]],
    source_rel: str,
    block_by_id: dict[str, dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for index, replacement in enumerate(replacements, start=1):
        if not isinstance(replacement, dict):
            continue
        item = dict(replacement)
        item.setdefault("replacement_id", f"X{index:03d}")
        block = (block_by_id or {}).get(str(item.get("block_id") or ""))
        if not str(item.get("source_relative_path") or "").strip():
            item["source_relative_path"] = str((block or {}).get("source_relative_path") or source_rel)
        if not str(item.get("location") or "").strip() and block:
            item["location"] = str(block.get("location") or "")
        out.append(item)
    return out


def finalize_replacements(replacements: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen_keys: set[str] = set()
    used_ids: set[str] = set()
    out: list[dict[str, Any]] = []
    for item in replacements:
        key = json.dumps(
            {
                "source": item.get("source_relative_path", ""),
                "block": item.get("block_id", ""),
                "old": item.get("old_text", ""),
                "new": item.get("new_text", ""),
            },
            ensure_ascii=False,
            sort_keys=True,
        )
        if key in seen_keys:
            continue
        seen_keys.add(key)
        item = dict(item)
        replacement_id = str(item.get("replacement_id") or "").strip()
        if not replacement_id or replacement_id in used_ids:
            replacement_id = f"X{len(out) + 1:04d}"
        item["replacement_id"] = replacement_id
        used_ids.add(replacement_id)
        out.append(item)
    return out


def dedupe_task_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    out: list[dict[str, Any]] = []
    for row in rows:
        key_payload = {
            "source": row.get("source_relative_path", ""),
            "block": row.get("block_id", ""),
            "quote": row.get("quote", ""),
            "result": row.get("result", ""),
            "values": row.get("values", {}),
        }
        key = json.dumps(key_payload, ensure_ascii=False, sort_keys=True)
        if key in seen:
            continue
        seen.add(key)
        out.append(row)
    return out


def call_task_llm(
    provider: str,
    model: str,
    task_instruction: str,
    user_query: str,
    block_map: Dict[str, Any],
    *,
    chunk_tokens: int,
    overlap_tokens: int,
    min_chunks: int,
    max_retries: int,
    max_output_tokens: int,
    timeout_sec: float,
    service_tier: str,
    workers: int = DEFAULT_WORKERS,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], Dict[str, Any]]:
    system_prompt = read_text(SYSTEM_PROMPT_FILE)
    blocks = [dict(block) for block in block_map.get("blocks", []) if str(block.get("text") or "").strip()]
    pairs = build_chunks(
        blocks,
        chunk_tokens=chunk_tokens,
        overlap_tokens=overlap_tokens,
        min_chunks=min_chunks,
        max_input_tokens=0,
        prompt_overhead_tokens=2500,
    )
    print(f"[DOC TASK] provider={provider} model={model} chunks={len(pairs)}")

    rows: list[dict[str, Any]] = []
    replacements: list[dict[str, Any]] = []
    usage_total = {"calls": 0, "input_tokens": 0, "output_tokens": 0, "total_tokens": 0, "reasoning_tokens": 0}
    source_rel = str(block_map.get("source_relative_path") or "")
    block_by_id = block_lookup(block_map)

    settings = load_settings()
    if provider == "openai":
        from openai import OpenAI
        from providers.openai_provider import call_json_object

        api_key = resolve_api_key("openai", settings)
        if not api_key:
            raise RuntimeError("OpenAI API key not found. Set OPENAI_API_KEY or config\api_key_openai.txt.")
        client = OpenAI(api_key=api_key, timeout=timeout_sec, max_retries=0)
        label = "OpenAI"

        def call(index: int, prompt: str):
            return call_json_object(
                client,
                model=model,
                instructions=system_prompt,
                user_prompt=prompt,
                reasoning_effort="low",
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max_retries,
                service_tier=service_tier,
                use_idempotency=True,
                doc_hash=hashlib.sha256(source_rel.encode("utf-8")).hexdigest(),
                chunk_index=index,
            )

    elif provider == "gemini":
        from google import genai
        from providers.gemini_provider import call_structured

        api_key = resolve_api_key("gemini", settings)
        if not api_key:
            raise RuntimeError("Gemini API key not found. Set GEMINI_API_KEY or config\api_key_gemini.txt.")
        client = genai.Client(api_key=api_key)
        label = "Gemini"

        def call(index: int, prompt: str):
            return call_structured(
                client,
                model=model,
                system_instruction=system_prompt,
                user_prompt=prompt,
                max_retries=max_retries,
            )

    elif provider == "xai":
        from providers.xai_provider import call_json_object

        api_key = resolve_api_key("xai", settings)
        if not api_key:
            raise RuntimeError("xAI API key not found. Set XAI_API_KEY or config\api_key_xai.txt.")
        label = "xAI"

        def call(index: int, prompt: str):
            return call_json_object(
                api_key=api_key,
                model=model,
                instructions=system_prompt,
                user_prompt=prompt,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max_retries,
                use_idempotency=True,
                doc_hash=hashlib.sha256(source_rel.encode("utf-8")).hexdigest(),
                chunk_index=index,
            )

    elif provider == "anthropic":
        from providers.anthropic_provider import call_json_object

        api_key = resolve_api_key("anthropic", settings)
        if not api_key:
            raise RuntimeError("Anthropic API key not found. Set ANTHROPIC_API_KEY or config\api_key_anthropic.txt.")
        label = "Claude"

        def call(index: int, prompt: str):
            return call_json_object(
                api_key=api_key,
                model=model,
                instructions=system_prompt,
                user_prompt=prompt,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                max_retries=max_retries,
            )

    else:
        raise RuntimeError(f"Unsupported document task provider: {provider}")

    # Chunks are independent, so results are collected per index and merged in order:
    # concurrent completion must not shuffle rows or replacements between runs.
    results: Dict[int, Dict[str, Any]] = {}
    lock = threading.Lock()
    total = len(pairs)

    def work(job: tuple[int, str, str]) -> None:
        index, overlap, chunk = job
        started = time.perf_counter()
        print(f"  [DOC TASK {index}/{total}] {label} start")
        prompt = build_prompt(task_instruction, user_query, block_map, overlap, chunk)
        obj, usage, tier = call(index, prompt)
        chunk_rows = normalize_rows(obj.get("rows") or [], source_rel, block_by_id)
        chunk_replacements = normalize_replacements(obj.get("replacements") or [], source_rel, block_by_id)
        with lock:
            results[index] = {"rows": chunk_rows, "replacements": chunk_replacements, "usage": usage, "tier": tier}
        print(
            f"  [DOC TASK {index}/{total}] {label} ok "
            f"time={format_duration(time.perf_counter() - started)} "
            f"rows={len(chunk_rows)} replacements={len(chunk_replacements)}"
        )

    jobs = [(index, overlap, chunk) for index, (overlap, chunk) in enumerate(pairs, start=1)]
    lanes = min(resolve_workers(workers), max(1, len(jobs)))
    if lanes > 1:
        print(f"[DOC TASK] {label} concurrency={lanes} chunks at a time")
    run_jobs(jobs, work, workers=lanes)

    for index in sorted(results):
        item = results[index]
        usage = item["usage"]
        usage_total["calls"] += 1
        for key in ("input_tokens", "output_tokens", "total_tokens", "reasoning_tokens"):
            usage_total[key] += int(usage.get(key, 0) or 0)
        if item["tier"]:
            usage_total["service_tier"] = item["tier"]
        rows.extend(item["rows"])
        replacements.extend(item["replacements"])

    return dedupe_task_rows(rows), finalize_replacements(replacements), usage_total


def apply_docx_replacements(source: Path, out_path: Path, replacements: list[dict[str, Any]]) -> dict[str, Any]:
    if source.suffix.lower() != ".docx":
        return {"status": "skipped", "reason": "exact replacements are implemented for DOCX only", "applied": 0, "skipped": len(replacements)}
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, out_path)

    refs_payload = block_paragraph_refs_docx(out_path)
    doc = refs_payload["document"]
    refs = refs_payload["refs"]
    applied = 0
    skipped: list[dict[str, Any]] = []
    for item in replacements:
        block_id = str(item.get("block_id") or "").strip()
        old_text = str(item.get("old_text") or "").strip()
        new_text = str(item.get("new_text") or "")
        paragraph = refs.get(block_id)
        if paragraph is None or not old_text:
            skipped.append({**item, "status": "skipped", "reason": "block or old_text missing"})
            continue
        current = paragraph.text
        if old_text not in current:
            skipped.append({**item, "status": "skipped", "reason": "old_text was not found exactly"})
            continue
        paragraph.text = current.replace(old_text, new_text, 1)
        applied += 1
    doc.save(str(out_path))
    return {"status": "ok", "edited": str(out_path), "applied": applied, "skipped": len(skipped), "skipped_items": skipped}


CONFIDENCE_RANKS = {
    "high": 3,
    "высокая": 3,
    "высокий": 3,
    "safe": 3,
    "medium": 2,
    "средняя": 2,
    "средний": 2,
    "low": 1,
    "низкая": 1,
    "низкий": 1,
}


def confidence_rank(value: Any) -> int:
    text = str(value or "").strip().lower().replace("_", " ").replace("-", " ")
    if not text:
        return 0
    for token, rank in CONFIDENCE_RANKS.items():
        if token in text:
            return rank
    return 0


def confidence_label(rank: int) -> str:
    if rank >= 3:
        return "high"
    if rank == 2:
        return "medium"
    if rank == 1:
        return "low"
    return "unknown"


def replacement_apply_decision(item: dict[str, Any], min_confidence: str = "high") -> tuple[bool, str]:
    action = str(item.get("action") or item.get("decision") or item.get("fix_mode") or "").strip().lower()
    if action in {"skip", "keep", "review", "manual", "manual_review", "unresolved", "ambiguous"}:
        return False, f"replacement action is {action}"
    if not str(item.get("block_id") or "").strip():
        return False, "block_id missing"
    if not str(item.get("old_text") or "").strip():
        return False, "old_text missing"
    threshold = confidence_rank(min_confidence) or 3
    rank = confidence_rank(item.get("confidence"))
    if rank < threshold:
        return False, f"confidence {confidence_label(rank)} is below {confidence_label(threshold)}"
    return True, ""


def apply_docx_replacements_safely(
    source: Path,
    out_path: Path,
    replacements: list[dict[str, Any]],
    *,
    min_confidence: str = "high",
    create_copy: bool = False,
) -> dict[str, Any]:
    safe_replacements: list[dict[str, Any]] = []
    for item in replacements:
        should_apply, reason = replacement_apply_decision(item, min_confidence)
        if should_apply:
            safe_replacements.append(item)
            continue
        item["status"] = "unresolved"
        item["status_note"] = reason

    if not safe_replacements and not create_copy:
        return {"status": "no_safe_replacements", "applied": 0, "skipped": len(replacements), "skipped_items": []}

    edit_payload = apply_docx_replacements(source, out_path, safe_replacements)
    if edit_payload.get("status") == "skipped":
        reason = str(edit_payload.get("reason") or "edit was skipped")
        for item in safe_replacements:
            item["status"] = "unresolved"
            item["status_note"] = reason
        return edit_payload

    skipped_by_id = {
        str(item.get("replacement_id") or ""): item
        for item in edit_payload.get("skipped_items", []) or []
        if isinstance(item, dict)
    }
    for item in safe_replacements:
        skipped = skipped_by_id.get(str(item.get("replacement_id") or ""))
        if skipped:
            item["status"] = "unresolved"
            item["status_note"] = str(skipped.get("reason") or "replacement was skipped")
        else:
            item["status"] = "applied"
            item["status_note"] = ""
    return edit_payload


BASE_ROW_KEYS = {"row_id", "source_relative_path", "block_id", "location", "quote", "result", "notes", "values"}


UNRESOLVED_ROW_MARKERS = {
    "unresolved",
    "manual_review",
    "needs_review",
    "ambiguous",
    "not_fixable",
    "not_applied",
    "review",
    "неразреш",
    "неоднознач",
    "сомн",
    "вручн",
}


def display_value(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value or "")


def structured_columns(rows: list[dict[str, Any]]) -> list[str]:
    columns: list[str] = []

    def add(name: Any) -> None:
        text = str(name or "").strip()
        if text and text not in columns:
            columns.append(text)

    for row in rows:
        values = row.get("values")
        if isinstance(values, dict):
            for key in values:
                add(key)
        for key in row:
            if key not in BASE_ROW_KEYS:
                add(key)
    return columns


def row_structured_value(row: dict[str, Any], column: str) -> str:
    values = row.get("values")
    if isinstance(values, dict) and column in values:
        return display_value(values.get(column))
    return display_value(row.get(column, ""))


def structured_details(row: dict[str, Any]) -> str:
    lines: list[str] = []
    for column in structured_columns([row]):
        value = row_structured_value(row, column).strip()
        if value:
            lines.append(f"{column}: {value}")
    return "\n".join(lines)


def row_unresolved_reason(row: dict[str, Any]) -> str:
    values = row.get("values") if isinstance(row.get("values"), dict) else {}
    parts: list[str] = []
    for key in ("status", "resolution", "action", "fix_mode", "result", "notes"):
        parts.append(str(row.get(key) or ""))
        parts.append(str(values.get(key) or ""))
    text = " ".join(parts).lower()
    for marker in UNRESOLVED_ROW_MARKERS:
        if marker in text:
            return str(row.get("notes") or values.get("reason") or values.get("resolution") or row.get("result") or "manual review")
    return ""


def build_unresolved_items(
    rows: list[dict[str, Any]],
    replacements: list[dict[str, Any]],
    *,
    include_row_findings: bool = False,
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for item in replacements:
        if str(item.get("status") or "").strip().lower() == "applied":
            continue
        out.append(
            {
                "kind": "replacement",
                "source_relative_path": item.get("source_relative_path", ""),
                "block_id": item.get("block_id", ""),
                "location": item.get("location", ""),
                "quote": item.get("old_text", ""),
                "proposed": item.get("new_text", ""),
                "reason": item.get("status_note") or item.get("reason", ""),
                "confidence": item.get("confidence", ""),
                "status": item.get("status", "unresolved"),
            }
        )
    if include_row_findings:
        for row in rows:
            reason = row_unresolved_reason(row)
            if not reason:
                continue
            out.append(
                {
                    "kind": "audit",
                    "source_relative_path": row.get("source_relative_path", ""),
                    "block_id": row.get("block_id", ""),
                    "location": row.get("location", ""),
                    "quote": row.get("quote", ""),
                    "proposed": row.get("result", ""),
                    "reason": reason,
                    "confidence": row.get("confidence", ""),
                    "status": "unresolved",
                }
            )
    return out


def count_replacements(replacements: list[dict[str, Any]], status: str) -> int:
    expected = str(status or "").strip().lower()
    return sum(1 for item in replacements if str(item.get("status") or "").strip().lower() == expected)


def write_task_xlsx(out_path: Path, payload: Dict[str, Any]) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    out_path.parent.mkdir(parents=True, exist_ok=True)
    rows = payload.get("rows", []) or []
    dynamic_columns = structured_columns(rows)

    wb = Workbook()
    ws = wb.active
    ws.title = "Results"
    base_headers = ["File", "Block ID", "Location", "Quote", "Result", "Notes"]
    ws.append(base_headers + dynamic_columns)
    for row in rows:
        ws.append(
            [
                row.get("source_relative_path", ""),
                row.get("block_id", ""),
                row.get("location", ""),
                row.get("quote", ""),
                row.get("result", ""),
                row.get("notes", ""),
                *[row_structured_value(row, column) for column in dynamic_columns],
            ]
        )

    ws_docs = wb.create_sheet("Documents")
    ws_docs.append(["File", "Type", "Blocks"])
    for item in payload.get("document_items", []) or []:
        ws_docs.append([
            item.get("source_relative_path", ""),
            item.get("document_type", ""),
            item.get("blocks", ""),
        ])

    ws2 = wb.create_sheet("Replacements")
    ws2.append(["File", "Block ID", "Location", "Old text", "New text", "Reason", "Confidence", "Status", "Status note"])
    for item in payload.get("replacements", []) or []:
        ws2.append([
            item.get("source_relative_path", ""),
            item.get("block_id", ""),
            item.get("location", ""),
            item.get("old_text", ""),
            item.get("new_text", ""),
            item.get("reason", ""),
            item.get("confidence", ""),
            item.get("status", ""),
            item.get("status_note", ""),
        ])

    ws3 = wb.create_sheet("Unresolved")
    ws3.append(["Kind", "File", "Block ID", "Location", "Quote", "Proposed", "Reason", "Confidence", "Status"])
    for item in payload.get("unresolved_items", []) or []:
        ws3.append([
            item.get("kind", ""),
            item.get("source_relative_path", ""),
            item.get("block_id", ""),
            item.get("location", ""),
            item.get("quote", ""),
            item.get("proposed", ""),
            item.get("reason", ""),
            item.get("confidence", ""),
            item.get("status", ""),
        ])

    for sheet in [ws, ws_docs, ws2, ws3]:
        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(name="Tahoma", size=9, bold=True, color="FFFFFF")
        text_font = Font(name="Tahoma", size=9, color="1F1F1F")
        even_fill = PatternFill("solid", fgColor="F7F3EA")
        odd_fill = PatternFill("solid", fgColor="EFE7D8")
        thin = Side(style="thin", color="D9E2F3")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        for cell in sheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = border
                if cell.row == 1:
                    continue
                cell.font = text_font
                cell.fill = even_fill if cell.row % 2 == 0 else odd_fill
                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        for column in sheet.columns:
            letter = column[0].column_letter
            max_len = max(len(str(cell.value or "")) for cell in column)
            sheet.column_dimensions[letter].width = min(max(max_len + 2, 10), 60)
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        sheet.sheet_view.showGridLines = False

    for letter in ["D", "E", "F"]:
        ws.column_dimensions[letter].width = max(ws.column_dimensions[letter].width or 0, 42)
    wb.save(str(out_path))


def compact_cell_text(value: Any) -> str:
    return " ".join(str(value or "").replace("\xa0", " ").replace("\r", " ").replace("\n", " ").split())


def unique_headers(headers: list[str]) -> list[str]:
    out: list[str] = []
    seen: dict[str, int] = {}
    for index, header in enumerate(headers, start=1):
        base = compact_cell_text(header) or f"Column {index}"
        count = seen.get(base, 0) + 1
        seen[base] = count
        out.append(base if count == 1 else f"{base} {count}")
    return out


def read_docx_table_template(path: Path) -> dict[str, Any]:
    from docx import Document

    doc = Document(str(path))
    if not doc.tables:
        raise RuntimeError(f"DOCX table template has no tables: {path}")
    table = doc.tables[0]
    if not table.rows:
        raise RuntimeError(f"DOCX table template has an empty first table: {path}")
    headers = unique_headers([cell.text for cell in table.rows[0].cells])
    sample = [compact_cell_text(cell.text) for cell in table.rows[1].cells] if len(table.rows) > 1 else []
    return {"path": str(path), "kind": "docx", "headers": headers, "sample": sample}


def read_xlsx_table_template(path: Path) -> dict[str, Any]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        ws = wb.worksheets[0]
        header_values: list[str] = []
        sample_values: list[str] = []
        header_row_index = 0
        for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [compact_cell_text(value) for value in row]
            if any(values):
                header_values = values
                header_row_index = row_index
                break
        if not header_values:
            raise RuntimeError(f"XLSX table template has no non-empty rows: {path}")
        for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if row_index <= header_row_index:
                continue
            values = [compact_cell_text(value) for value in row]
            if any(values):
                sample_values = values
                break
        return {"path": str(path), "kind": "xlsx", "headers": unique_headers(header_values), "sample": sample_values}
    finally:
        wb.close()


def read_table_template(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_table_template(path)
    if suffix == ".xlsx":
        return read_xlsx_table_template(path)
    raise RuntimeError(f"Clean table template must be DOCX or XLSX: {path}")


def resolve_clean_template(input_dir: Path, sources: list[Path], template_ref: str = "") -> Path | None:
    text = str(template_ref or "").strip()
    if text:
        candidate = Path(text)
        if not candidate.is_absolute():
            candidate = input_dir / text
        candidate = candidate.resolve()
        if not candidate.exists() or not candidate.is_file():
            raise RuntimeError(f"Clean table template was not found: {candidate}")
        if candidate.suffix.lower() not in {".docx", ".xlsx"}:
            raise RuntimeError("Clean table template must be a DOCX or XLSX file.")
        return candidate

    candidates = [source for source in sources if source.suffix.lower() in {".docx", ".xlsx"}]
    if len(candidates) == 1 and len(sources) > 1:
        return candidates[0]
    return None


def normalize_clean_table_value(header: str, value: Any, sample: str = "") -> str:
    text = compact_cell_text(value)
    sample_text = compact_cell_text(sample)
    if not text:
        return ""
    if sample_text.startswith("от ") and re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{4}", text):
        text = f"от {text}"
    if sample_text.startswith("№") or header.strip().lower() in {"номер", "№"}:
        match = re.search(r"(\d+)", text)
        if match:
            text = f"№ {match.group(1)}"
    text = re.sub(r"№\s*", "№ ", text)
    return compact_cell_text(text)


def is_clean_date_column(header: str, sample: str = "") -> bool:
    return compact_cell_text(header).lower() in {"дата", "date"} or compact_cell_text(sample).startswith("от ")


def is_clean_number_column(header: str, sample: str = "") -> bool:
    header_text = compact_cell_text(header).lower()
    return header_text in {"номер", "№", "number", "num"} or compact_cell_text(sample).startswith("№")


def strip_clean_reference_tail(value: str, *, strip_date: bool, strip_number: bool) -> str:
    text = compact_cell_text(value)
    if not text:
        return ""
    if strip_date:
        text = re.sub(
            r"\s*,?\s*от\s+\d{1,2}\.\d{1,2}\.\d{4}\s*(?:г\.?)?\s*(?:[,;]?\s*№\s*\d+[\w/-]*)?\s*$",
            "",
            text,
            flags=re.IGNORECASE,
        )
    if strip_number:
        text = re.sub(r"\s*,?\s*№\s*\d+[\w/-]*\s*$", "", text, flags=re.IGNORECASE)
    return compact_cell_text(text)


def clean_table_rows(rows: list[dict[str, Any]], headers: list[str], sample: list[str] | None = None) -> list[list[str]]:
    sample = sample or []
    date_columns = {
        index for index, header in enumerate(headers) if is_clean_date_column(header, sample[index] if index < len(sample) else "")
    }
    number_columns = {
        index for index, header in enumerate(headers) if is_clean_number_column(header, sample[index] if index < len(sample) else "")
    }
    has_date_column = bool(date_columns)
    has_number_column = bool(number_columns)
    out: list[list[str]] = []
    for row in rows:
        values = row.get("values") if isinstance(row.get("values"), dict) else {}
        if not isinstance(values, dict):
            values = {}
        item: list[str] = []
        for index, header in enumerate(headers):
            raw_value = values.get(header, row.get(header, ""))
            text = normalize_clean_table_value(header, raw_value, sample[index] if index < len(sample) else "")
            if index not in date_columns and index not in number_columns:
                text = strip_clean_reference_tail(text, strip_date=has_date_column, strip_number=has_number_column)
            item.append(text)
        if any(cell.strip() for cell in item):
            out.append(item)
    return out


def write_clean_table_xlsx(out_path: Path, headers: list[str], rows: list[list[str]]) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Реквизиты"
    ws.append(headers)
    for row in rows:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    header_font = Font(name="Tahoma", size=10, bold=True, color="1F1F1F")
    body_font = Font(name="Tahoma", size=10, color="1F1F1F")
    thin = Side(style="thin", color="808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for worksheet_row in ws.iter_rows():
        for cell in worksheet_row:
            cell.border = border
            cell.font = body_font
            cell.alignment = Alignment(wrap_text=True, vertical="center")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for column_index, header in enumerate(headers, start=1):
        max_len = max(len(str(ws.cell(row_index, column_index).value or "")) for row_index in range(1, ws.max_row + 1))
        width = min(max(max_len + 4, len(header) + 2, 12), 64)
        ws.column_dimensions[get_column_letter(column_index)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.sheet_view.showGridLines = False
    wb.save(str(out_path))


def set_docx_cell_text(cell: Any, text: str, *, bold: bool = False, center: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    cell.text = str(text or "")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Tahoma"
            run.font.size = Pt(9)
            run.font.bold = bold


def write_filled_template_docx(template_path: Path, out_path: Path, headers: list[str], rows: list[list[str]]) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template_path))
    if not doc.tables:
        raise RuntimeError(f"DOCX table template has no tables: {template_path}")
    table = doc.tables[0]
    if len(table.columns) != len(headers):
        raise RuntimeError(
            f"DOCX table template column count ({len(table.columns)}) does not match clean table headers ({len(headers)})."
        )
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    for column_index, header in enumerate(headers):
        set_docx_cell_text(table.cell(0, column_index), header, bold=True, center=True)
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            center = column_index >= len(headers) - 2
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_docx_cell_text(cell, value, center=center)
    doc.save(str(out_path))


def write_clean_table_outputs(report_dir: Path, stamp: str, payload: Dict[str, Any], template_path: Path) -> dict[str, str]:
    template = read_table_template(template_path)
    headers = [str(header) for header in template.get("headers", []) if str(header).strip()]
    if not headers:
        raise RuntimeError(f"Clean table template has no headers: {template_path}")
    rows = clean_table_rows(payload.get("rows", []) or [], headers, template.get("sample", []) or [])
    clean_payload = {
        "status": "ok",
        "template_path": str(template_path),
        "headers": headers,
        "rows": rows,
        "source_rows": len(payload.get("rows", []) or []),
    }
    outputs: dict[str, str] = {}
    json_path = report_dir / f"{stamp}__doc_task_clean.json"
    xlsx_path = report_dir / f"{stamp}__doc_task_clean.xlsx"
    write_json(json_path, clean_payload)
    write_clean_table_xlsx(xlsx_path, headers, rows)
    outputs["json"] = str(json_path)
    outputs["xlsx"] = str(xlsx_path)
    if template_path.suffix.lower() == ".docx":
        docx_path = report_dir / f"{stamp}__doc_task_clean.docx"
        write_filled_template_docx(template_path, docx_path, headers, rows)
        outputs["docx"] = str(docx_path)
    return outputs


def write_task_docx(out_path: Path, payload: Dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    out_path.parent.mkdir(parents=True, exist_ok=True)

    def shade_cell(cell: Any, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    def set_cell_font(cell: Any, *, bold: bool = False, color: str = "1F1F1F") -> None:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = "Tahoma"
                run.font.size = Pt(9)
                run.font.bold = bold
                run.font.color.rgb = RGBColor.from_string(color)

    def style_table(table: Any) -> None:
        table.style = "Table Grid"
        for row_index, row in enumerate(table.rows):
            fill = "1F4E78" if row_index == 0 else ("F7F3EA" if row_index % 2 == 1 else "EFE7D8")
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                shade_cell(cell, fill)
                set_cell_font(cell, bold=row_index == 0, color="FFFFFF" if row_index == 0 else "1F1F1F")

    def add_meta_row(table: Any, label: str, value: Any) -> None:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value or "")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style = doc.styles["Normal"]
    style.font.name = "Tahoma"
    style.font.size = Pt(10)

    title = doc.add_heading("Audion Docs AI - Document Task Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = doc.add_table(rows=0, cols=2)
    add_meta_row(meta, "Task", payload.get("task_label", ""))
    add_meta_row(meta, "Provider/model", f"{payload.get('provider', '')} / {payload.get('model', '')}")
    add_meta_row(meta, "Scope", payload.get("task_scope", ""))
    add_meta_row(meta, "Documents", payload.get("documents", 0))
    add_meta_row(meta, "Rows", len(payload.get("rows", []) or []))
    add_meta_row(meta, "Replacements", len(payload.get("replacements", []) or []))
    add_meta_row(meta, "Applied replacements", count_replacements(payload.get("replacements", []) or [], "applied"))
    add_meta_row(meta, "Unresolved items", len(payload.get("unresolved_items", []) or []))
    style_table(meta)

    summary = str(payload.get("summary") or "").strip()
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    rows = payload.get("rows", []) or []
    if rows:
        doc.add_heading("Results", level=1)
        table = doc.add_table(rows=1, cols=5)
        headers = ["File / block", "Quote", "Result", "Details", "Notes"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for row in rows:
            cells = table.add_row().cells
            location = str(row.get("location") or "").strip()
            where = f"{row.get('source_relative_path', '')}\n{row.get('block_id', '')}"
            if location:
                where = f"{where}\n{location}"
            cells[0].text = where
            cells[1].text = str(row.get("quote", ""))
            cells[2].text = str(row.get("result", ""))
            cells[3].text = structured_details(row)
            cells[4].text = str(row.get("notes", ""))
        style_table(table)

    replacements = payload.get("replacements", []) or []
    if replacements:
        doc.add_heading("Exact Replacements", level=1)
        table = doc.add_table(rows=1, cols=7)
        headers = ["File", "Block", "Old text", "New text", "Reason", "Confidence", "Status"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for item in replacements:
            cells = table.add_row().cells
            cells[0].text = str(item.get("source_relative_path", ""))
            cells[1].text = str(item.get("block_id", ""))
            cells[2].text = str(item.get("old_text", ""))
            cells[3].text = str(item.get("new_text", ""))
            cells[4].text = str(item.get("reason", ""))
            cells[5].text = str(item.get("confidence", ""))
            cells[6].text = str(item.get("status", ""))
        style_table(table)

    unresolved = payload.get("unresolved_items", []) or []
    if unresolved:
        doc.add_heading("Unresolved Items", level=1)
        table = doc.add_table(rows=1, cols=7)
        headers = ["Kind", "File / block", "Quote", "Proposed", "Reason", "Confidence", "Status"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for item in unresolved:
            cells = table.add_row().cells
            location = str(item.get("location") or "").strip()
            where = f"{item.get('source_relative_path', '')}\n{item.get('block_id', '')}"
            if location:
                where = f"{where}\n{location}"
            cells[0].text = str(item.get("kind", ""))
            cells[1].text = where
            cells[2].text = str(item.get("quote", ""))
            cells[3].text = str(item.get("proposed", ""))
            cells[4].text = str(item.get("reason", ""))
            cells[5].text = str(item.get("confidence", ""))
            cells[6].text = str(item.get("status", ""))
        style_table(table)
    doc.save(str(out_path))


def write_task_markdown(out_path: Path, payload: Dict[str, Any]) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# Audion Docs AI - Document Task Report",
        "",
        f"- Task: {payload.get('task_label', '')}",
        f"- Provider/model: {payload.get('provider', '')} / {payload.get('model', '')}",
        f"- Documents: {payload.get('documents', 0)}",
        f"- Rows: {len(payload.get('rows', []) or [])}",
        f"- Replacements: {len(payload.get('replacements', []) or [])}",
        f"- Applied replacements: {count_replacements(payload.get('replacements', []) or [], 'applied')}",
        f"- Unresolved items: {len(payload.get('unresolved_items', []) or [])}",
        "",
    ]
    summary = str(payload.get("summary") or "").strip()
    if summary:
        lines.extend(["## Summary", "", summary, ""])
    if payload.get("rows"):
        lines.extend(["## Results", ""])
        for row in payload["rows"]:
            details = display_value(row.get("values", {}))
            suffix = f" | {details}" if details and details != "{}" else ""
            lines.append(f"- `{row.get('source_relative_path', '')}` `{row.get('block_id', '')}`: {row.get('result', '')}{suffix}")
        lines.append("")
    if payload.get("replacements"):
        lines.extend(["## Exact Replacements", ""])
        for item in payload["replacements"]:
            lines.append(f"- `{item.get('source_relative_path', '')}` `{item.get('block_id', '')}`: `{item.get('old_text', '')}` -> `{item.get('new_text', '')}` ({item.get('status', '')})")
        lines.append("")
    if payload.get("unresolved_items"):
        lines.extend(["## Unresolved Items", ""])
        for item in payload["unresolved_items"]:
            lines.append(
                f"- `{item.get('source_relative_path', '')}` `{item.get('block_id', '')}` "
                f"[{item.get('kind', '')}]: {item.get('reason', '')}"
            )
        lines.append("")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def run_doc_task(
    paths,
    *,
    recursive: bool,
    task_scope: str,
    provider: str,
    model: str,
    task_ref: str,
    user_query: str,
    apply_replacements: bool,
    apply_confidence: str,
    docx_only: bool,
    pdf_max_pages: int,
    clean_table: bool,
    clean_table_template: str,
    chunk_tokens: int,
    overlap_tokens: int,
    min_chunks: int,
    max_retries: int,
    max_output_tokens: int,
    timeout_sec: float,
    service_tier: str,
    resume: bool,
    workers: int = DEFAULT_WORKERS,
) -> Dict[str, Any]:
    ensure_project_dirs(paths)
    settings = load_settings()
    if not model:
        if provider == "openai":
            model = resolve_model("openai", "audit", settings)
        elif provider == "gemini":
            model = resolve_model("gemini", "audit_fast", settings)
        elif provider == "xai":
            model = resolve_model("xai", "audit", settings)
        elif provider == "anthropic":
            model = resolve_model("anthropic", "audit", settings)
    if not model:
        raise RuntimeError(
            f"No model selected for {provider}. Choose a model in the GUI dropdown, "
            "or add/check a favorite model so config\\gui_model_cache.json can provide the fallback."
        )
    task_instruction, task_entry = load_doc_task(task_ref)
    task_ref_resolved = str(task_entry.get("ref") or task_ref)
    system_prompt_sha256 = sha256_file(SYSTEM_PROMPT_FILE) if SYSTEM_PROMPT_FILE.exists() else ""
    if bool(task_entry.get("quick")):
        remember_quick_doc_task_use(str(task_entry.get("quick_ref") or "").strip())
    else:
        remember_doc_task_use(task_ref_resolved)

    sources = iter_task_inputs(paths.input_dir, recursive=recursive)
    if docx_only:
        sources = [source for source in sources if source.suffix.lower() == ".docx"]
    clean_template_path = resolve_clean_template(paths.input_dir, sources, clean_table_template) if clean_table else None
    if clean_template_path:
        sources = [source for source in sources if source.resolve() != clean_template_path]
        print(f"[CLEAN TEMPLATE] {clean_template_path}")
    elif clean_table:
        print("[CLEAN TEMPLATE] no single DOCX/XLSX template found; clean table output will be skipped.")
    if not sources:
        expected = ".docx files" if docx_only else ".docx/.pptx/.xlsx/.pdf files"
        print(f"[INFO] No {expected} found in: {paths.input_dir}")
        return {"status": "empty", "documents": 0, "rows": [], "replacements": []}

    scope = resolve_task_scope(task_scope, len(sources), apply_replacements)
    print(f"[DOC TASK] inputs={len(sources)} scope={scope} recursive={recursive}")

    aggregate_rows: list[dict[str, Any]] = []
    aggregate_replacements: list[dict[str, Any]] = []
    usage_total = {"calls": 0, "input_tokens": 0, "output_tokens": 0, "total_tokens": 0, "reasoning_tokens": 0}
    edits: list[dict[str, Any]] = []
    document_items: list[dict[str, Any]] = []

    if scope == "corpus":
        block_maps: list[dict[str, Any]] = []
        for index, source in enumerate(sources, start=1):
            rel = source_relative(source, paths.input_dir).as_posix()
            print(f"[DOC TASK MAP] [{index}/{len(sources)}] {rel}")
            outputs = task_artifact_paths(paths, source, provider, model, edited_suffix="__doc_task")
            block_map = build_task_block_map(source, paths.input_dir, pdf_max_pages=pdf_max_pages)
            write_json(outputs["block_map"], block_map)
            block_maps.append(block_map)
            document_items.append(document_item_from_block_map(block_map))

        corpus_map = build_corpus_block_map(block_maps)
        cache_id = source_inventory_hash(
            sources,
            paths.input_dir,
            task_ref_resolved,
            sha256_text(task_instruction),
            user_query,
            provider,
            model,
            str(chunk_tokens),
            str(overlap_tokens),
            str(min_chunks),
            str(max_output_tokens),
            str(pdf_max_pages),
            system_prompt_sha256,
        )
        outputs = corpus_artifact_paths(paths, cache_id, provider, model)
        write_json(outputs["block_map"], corpus_map)

        if resume and outputs["cache"].exists():
            payload = read_json(outputs["cache"])
            rows = payload.get("rows", []) or []
            replacements = payload.get("replacements", []) or []
            usage = payload.get("usage", {}) or {}
            print(f"[DOC TASK] corpus cache=hit rows={len(rows)} replacements={len(replacements)}")
        else:
            rows, replacements, usage = call_task_llm(
                provider,
                model,
                task_instruction,
                user_query,
                corpus_map,
                chunk_tokens=chunk_tokens,
                overlap_tokens=overlap_tokens,
                min_chunks=min_chunks,
                max_retries=max_retries,
                max_output_tokens=max_output_tokens,
                timeout_sec=timeout_sec,
                service_tier=service_tier,
                workers=workers,
            )
            write_json(outputs["cache"], {"rows": rows, "replacements": replacements, "usage": usage})

        for key in usage_total:
            usage_total[key] += int(usage.get(key, 0) or 0)
        aggregate_rows.extend(rows)
        aggregate_replacements.extend(replacements)
        write_json(
            outputs["task"],
            {
                "status": "ok",
                "source_relative_path": "__corpus__",
                "provider": provider,
                "model": model,
                "task_ref": task_ref_resolved,
                "task_label": task_entry.get("label", ""),
                "rows": rows,
                "replacements": replacements,
                "usage": usage,
                "documents": document_items,
            },
        )
        print(f"[OK] corpus rows={len(rows)} replacements={len(replacements)}")
    else:
        for index, source in enumerate(sources, start=1):
            rel = source_relative(source, paths.input_dir).as_posix()
            print(f"[DOC TASK] [{index}/{len(sources)}] {rel}")
            outputs = task_artifact_paths(paths, source, provider, model, edited_suffix="__doc_task")
            block_map = build_task_block_map(source, paths.input_dir, pdf_max_pages=pdf_max_pages)
            write_json(outputs["block_map"], block_map)
            document_items.append(document_item_from_block_map(block_map))
            source_hash = sha256_file(source)
            cache_signature = doc_task_cache_signature(
                source_sha256=source_hash,
                task_ref=task_ref_resolved,
                task_instruction=task_instruction,
                user_query=user_query,
                provider=provider,
                model=model,
                chunk_tokens=chunk_tokens,
                overlap_tokens=overlap_tokens,
                min_chunks=min_chunks,
                max_output_tokens=max_output_tokens,
                pdf_max_pages=pdf_max_pages,
                system_prompt_sha256=system_prompt_sha256,
            )

            if resume and outputs["cache"].exists():
                payload = read_json(outputs["cache"])
                stale_reason = doc_task_cache_mismatch_reason(payload, cache_signature)
                if not stale_reason:
                    rows = payload.get("rows", []) or []
                    replacements = payload.get("replacements", []) or []
                    usage = payload.get("usage", {}) or {}
                    print(f"[DOC TASK] cache=hit rows={len(rows)} replacements={len(replacements)}")
                else:
                    print(f"[DOC TASK] cache=stale reason={stale_reason}; recomputing")
                    rows, replacements, usage = call_task_llm(
                        provider,
                        model,
                        task_instruction,
                        user_query,
                        block_map,
                        chunk_tokens=chunk_tokens,
                        overlap_tokens=overlap_tokens,
                        min_chunks=min_chunks,
                        max_retries=max_retries,
                        max_output_tokens=max_output_tokens,
                        timeout_sec=timeout_sec,
                        service_tier=service_tier,
                        workers=workers,
                    )
                    write_json(
                        outputs["cache"],
                        {
                            "source_sha256": source_hash,
                            "cache_signature": cache_signature,
                            "rows": rows,
                            "replacements": replacements,
                            "usage": usage,
                        },
                    )
            else:
                rows, replacements, usage = call_task_llm(
                    provider,
                    model,
                    task_instruction,
                    user_query,
                    block_map,
                    chunk_tokens=chunk_tokens,
                    overlap_tokens=overlap_tokens,
                    min_chunks=min_chunks,
                    max_retries=max_retries,
                    max_output_tokens=max_output_tokens,
                    timeout_sec=timeout_sec,
                    service_tier=service_tier,
                    workers=workers,
                )
                write_json(
                    outputs["cache"],
                    {
                        "source_sha256": source_hash,
                        "cache_signature": cache_signature,
                        "rows": rows,
                        "replacements": replacements,
                        "usage": usage,
                    },
                )

            for key in usage_total:
                usage_total[key] += int(usage.get(key, 0) or 0)
            aggregate_rows.extend(rows)
            aggregate_replacements.extend(replacements)

            edit_payload: dict[str, Any] = {"status": "not_requested", "applied": 0, "skipped": 0}
            if apply_replacements and replacements:
                edit_payload = apply_docx_replacements_safely(
                    source,
                    outputs["edited"],
                    replacements,
                    min_confidence=apply_confidence,
                    create_copy=False,
                )
                edits.append({"source_relative_path": rel, **edit_payload})

            doc_payload = {
                "status": "ok",
                "source_relative_path": rel,
                "provider": provider,
                "model": model,
                "task_ref": task_ref_resolved,
                "task_label": task_entry.get("label", ""),
                "rows": rows,
                "replacements": replacements,
                "usage": usage,
                "edit": edit_payload,
            }
            write_json(outputs["task"], doc_payload)
            print(f"[OK] rows={len(rows)} replacements={len(replacements)} edit={edit_payload.get('status')}")

    stamp = time.strftime("%Y-%m-%d_%H-%M-%S")
    report_dir = paths.output_dir / "_doc_tasks"
    aggregate = {
        "status": "ok",
        "created_at": stamp,
        "provider": provider,
        "model": model,
        "task_scope": scope,
        "task_ref": task_ref_resolved,
        "task_label": task_entry.get("label", ""),
        "task_note": task_entry.get("note", ""),
        "user_query": user_query,
        "apply_confidence": apply_confidence,
        "documents": len(sources),
        "document_items": document_items,
        "rows": aggregate_rows,
        "replacements": aggregate_replacements,
        "usage": usage_total,
        "edits": edits,
    }
    aggregate["unresolved_items"] = build_unresolved_items(
        aggregate_rows,
        aggregate_replacements,
        include_row_findings=False,
    )
    aggregate["summary"] = (
        f"Processed {len(sources)} input file(s) in {scope} scope, rows={len(aggregate_rows)}, "
        f"replacements={len(aggregate_replacements)}, "
        f"applied={count_replacements(aggregate_replacements, 'applied')}, "
        f"unresolved={len(aggregate['unresolved_items'])}."
    )
    report_stem = "doc_task"
    write_json(report_dir / f"{stamp}__{report_stem}.json", aggregate)
    write_task_xlsx(report_dir / f"{stamp}__{report_stem}.xlsx", aggregate)
    write_task_docx(report_dir / f"{stamp}__{report_stem}.docx", aggregate)
    write_task_markdown(report_dir / f"{stamp}__{report_stem}.md", aggregate)
    if clean_template_path:
        aggregate["clean_outputs"] = write_clean_table_outputs(report_dir, stamp, aggregate, clean_template_path)
        write_json(report_dir / f"{stamp}__{report_stem}.json", aggregate)
        for label, path in aggregate["clean_outputs"].items():
            print(f"[CLEAN {label.upper()}] {path}")
    print(f"[REPORT] {report_dir / f'{stamp}__{report_stem}.docx'}")
    return aggregate


def main() -> int:
    parser = argparse.ArgumentParser(description="Run a general Audion Docs AI document task.")
    sub = parser.add_subparsers(dest="cmd", required=True)
    p_run = sub.add_parser("run")
    p_run.add_argument("--recursive", action="store_true", default=True)
    p_run.add_argument("--no-recursive", dest="recursive", action="store_false")
    p_run.add_argument("--task-scope", choices=["auto", "document", "corpus"], default="auto")
    p_run.add_argument("--provider", choices=["openai", "gemini", "xai", "anthropic"], default="openai")
    p_run.add_argument("--model", default="")
    p_run.add_argument("--task-ref", default="")
    p_run.add_argument("--query", default="")
    p_run.add_argument("--apply-replacements", action="store_true")
    p_run.add_argument("--apply-confidence", choices=["high", "medium", "low"], default="high")
    p_run.add_argument("--docx-only", action="store_true")
    p_run.add_argument("--pdf-max-pages", type=int, default=5, help="Read only the first N pages from PDF inputs. Use 0 for all pages.")
    p_run.add_argument("--clean-table", action="store_true", help="Write clean table outputs from row values using a DOCX/XLSX template.")
    p_run.add_argument("--clean-table-template", default="", help="Template file path or input-relative path for clean table outputs.")
    p_run.add_argument("--chunk-tokens", type=int, default=12000)
    p_run.add_argument("--overlap-tokens", type=int, default=0)
    p_run.add_argument("--min-chunks", type=int, default=1)
    p_run.add_argument("--max-retries", type=int, default=1)
    p_run.add_argument("--max-output-tokens", type=int, default=32000)
    p_run.add_argument("--timeout-sec", type=float, default=240.0)
    p_run.add_argument("--service-tier", default="default")
    p_run.add_argument("--resume", action="store_true")
    p_run.add_argument("--workers", type=int, default=DEFAULT_WORKERS, help="chunks processed at a time")
    args = parser.parse_args()

    paths = default_project_paths(THIS_DIR.parent)
    query = args.query or os.environ.get(TASK_QUERY_ENV, "")
    if args.cmd == "run":
        run_doc_task(
            paths,
            recursive=args.recursive,
            task_scope=args.task_scope,
            provider=args.provider,
            model=args.model,
            task_ref=args.task_ref,
            user_query=query,
            apply_replacements=bool(args.apply_replacements),
            apply_confidence=args.apply_confidence,
            docx_only=bool(args.docx_only),
            pdf_max_pages=int(args.pdf_max_pages),
            clean_table=bool(args.clean_table),
            clean_table_template=str(args.clean_table_template or ""),
            chunk_tokens=args.chunk_tokens,
            overlap_tokens=args.overlap_tokens,
            min_chunks=args.min_chunks,
            max_retries=args.max_retries,
            max_output_tokens=args.max_output_tokens,
            timeout_sec=args.timeout_sec,
            service_tier=args.service_tier,
            resume=bool(args.resume),
            workers=int(args.workers),
        )
        return 0
    raise RuntimeError(f"Unknown command: {args.cmd}")


if __name__ == "__main__":
    raise SystemExit(main())
