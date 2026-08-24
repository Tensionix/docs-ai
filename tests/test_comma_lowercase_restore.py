from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
CORE = ROOT / "system_core"
if str(CORE) not in sys.path:
    sys.path.insert(0, str(CORE))

from comma_lowercase_restore import main as restore_main


def write_table_docx(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    doc.save(str(path))


def table_text(path: Path) -> str:
    doc = Document(str(path))
    return doc.tables[0].cell(0, 0).paragraphs[0].text


def write_report(path: Path, source: Path, before: str, after: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    hits = []
    for index, (word_before, word_after) in enumerate((("Ясенево", "ясенево"), ("Завершены", "завершены")), start=1):
        position = before.index(word_before)
        hits.append(
            {
                "kind": "table_cell",
                "table": "1",
                "row": 1,
                "cell": 1,
                "paragraph": 1,
                "hit": index,
                "position": position,
                "word_before": word_before,
                "word_after": word_after,
                "before": before,
                "after": after,
                "after_fragment": after[position:],
                "following_text": after[position + len(word_after) :],
            }
        )
    payload = {
        "command": "comma-lowercase-docx",
        "dry_run": False,
        "input": str(source.parent),
        "output": "",
        "scope": "table-cells",
        "files": 1,
        "changed_files": 1,
        "changes": len(hits),
        "results": [
            {
                "file": str(source),
                "output": "",
                "scope": "table-cells",
                "changes": len(hits),
                "locations": [],
                "hits": hits,
            }
        ],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


class CommaLowercaseRestoreTests(unittest.TestCase):
    def test_restore_map_restores_only_selected_hit(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_text = "г. Москва, Ясенево. Работы, Завершены."
            fixed_text = "г. Москва, ясенево. Работы, завершены."
            source = root / "input" / "sample.docx"
            fixed = root / "fixed" / "sample_comma_lowercase.docx"
            report = root / "report" / "comma_lowercase.json"
            restore_map = root / "config" / "comma_restore_map.yaml"
            output = root / "output"

            write_table_docx(source, source_text)
            write_table_docx(fixed, fixed_text)
            write_report(report, source, source_text, fixed_text)
            restore_map.parent.mkdir(parents=True, exist_ok=True)
            restore_map.write_text("restore_words:\n  ясенево: Ясенево\n", encoding="utf-8")

            rc = restore_main(
                [
                    "--report",
                    str(report),
                    "--source",
                    str(source),
                    "--fixed",
                    str(fixed),
                    "--output",
                    str(output),
                    "--restore-map",
                    str(restore_map),
                    "--json-out",
                    str(root / "reports_out" / "restore.json"),
                    "--md-out",
                    str(root / "reports_out" / "restore.md"),
                ]
            )

            self.assertEqual(rc, 0)
            restored = output / "sample_comma_lowercase_restored.docx"
            self.assertEqual(table_text(restored), "г. Москва, Ясенево. Работы, завершены.")

    def test_llm_decisions_json_can_restore_without_map(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_text = "г. Москва, Ясенево. Работы, Завершены."
            fixed_text = "г. Москва, ясенево. Работы, завершены."
            source = root / "input" / "sample.docx"
            fixed = root / "fixed" / "sample_comma_lowercase.docx"
            report = root / "report" / "comma_lowercase.json"
            decisions = root / "report" / "decisions.json"
            output = root / "output"

            write_table_docx(source, source_text)
            write_table_docx(fixed, fixed_text)
            write_report(report, source, source_text, fixed_text)
            decisions.write_text(
                json.dumps(
                    {
                        "decisions": [
                            {"file": "sample.docx", "hit": 1, "action": "restore", "reason": "топоним"},
                            {"file": "sample.docx", "hit": 2, "action": "keep", "reason": "обычное слово"},
                        ]
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            rc = restore_main(
                [
                    "--report",
                    str(report),
                    "--source",
                    str(source),
                    "--fixed",
                    str(fixed),
                    "--output",
                    str(output),
                    "--restore-map",
                    "",
                    "--decisions",
                    str(decisions),
                    "--json-out",
                    str(root / "reports_out" / "restore.json"),
                    "--md-out",
                    str(root / "reports_out" / "restore.md"),
                ]
            )

            self.assertEqual(rc, 0)
            restored = output / "sample_comma_lowercase_restored.docx"
            self.assertEqual(table_text(restored), "г. Москва, Ясенево. Работы, завершены.")


if __name__ == "__main__":
    unittest.main()
