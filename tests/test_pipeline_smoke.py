from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook, load_workbook

ROOT = Path(__file__).resolve().parents[1]
CORE = ROOT / "system_core"
if str(CORE) not in sys.path:
    sys.path.insert(0, str(CORE))

from document_model import build_block_map, default_project_paths, ensure_project_dirs, iter_documents
from doc_task_document_model import build_corpus_block_map, build_task_block_map, iter_task_inputs
import doc_task_resolver
from doc_task_resolver import TASK_TEXT_ENV, load_doc_task
from doc_task_runner import (
    apply_docx_replacements_safely,
    build_prompt,
    build_unresolved_items,
    doc_task_cache_mismatch_reason,
    doc_task_cache_signature,
    resolve_clean_template,
    source_inventory_hash,
    write_clean_table_outputs,
    write_task_docx,
    write_task_xlsx,
)
from document_normalizer import normalize_from_audit_logs
from config_resolver import load_settings, resolve_model
from pipeline import (
    artifact_paths,
    llm_cache_mismatch_reason,
    llm_cache_signature,
    invalid_human_report_fields,
    language_repair_provider_chain,
    normalize_issues,
    repair_human_report_language,
    validate_human_report_language,
    strip_error_anchors_from_document,
    write_annotated_document,
    write_audit_table,
)
from llm_audit_helpers import build_instructions
from render.render_map import build_human_location


def make_docx(path: Path, text: str = "Hello audit world") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(text)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell text"
    doc.save(str(path))


def make_xlsx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Objects"
    ws.append(["Name", "Address"])
    ws.append(["Central Clinic", "Omsk"])
    wb.save(str(path))


def make_pdf(path: Path) -> None:
    import pymupdf as fitz

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Project name: North Plant\nCustomer: ACME")
    page = doc.new_page()
    page.insert_text((72, 72), "This second page should be outside the first-page limit.")
    doc.save(str(path))
    doc.close()


def make_requisites_template_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    headers = ["Наименование", "Вид градостроительной документации", "Кем утверждено", "Дата", "номер"]
    sample = [
        "Ильинское сельское поселение",
        "Об утверждении правил землепользования и застройки Ильинского сельского поселения",
        "решение Совета Ильинского сельского поселения",
        "от 26.06.2024",
        "№ 94",
    ]
    doc = Document()
    table = doc.add_table(rows=2, cols=len(headers))
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    for col, value in enumerate(sample):
        table.cell(1, col).text = value
    doc.save(str(path))


class PipelineSmokeTests(unittest.TestCase):
    def test_application_model_defaults(self) -> None:
        settings = load_settings()
        self.assertEqual(resolve_model("openai", "audit", settings), "gpt-5.6-luna")
        self.assertEqual(resolve_model("gemini", "audit_fast", settings), "gemini-3.6-flash")
        self.assertEqual(resolve_model("xai", "audit", settings), "grok-4.3")

    def test_scan_ignores_office_temp_and_keeps_duplicate_stems(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            make_docx(paths.input_dir / "A" / "Report.docx")
            make_docx(paths.input_dir / "B" / "Report.docx")
            make_docx(paths.input_dir / "~$Temp.docx")

            docs = [p.relative_to(paths.input_dir).as_posix() for p in iter_documents(paths.input_dir)]

            self.assertEqual(docs, ["A/Report.docx", "B/Report.docx"])
            self.assertEqual(artifact_paths(paths, paths.input_dir / "A" / "Report.docx")["table"].parent, paths.output_dir / "A")
            self.assertEqual(artifact_paths(paths, paths.input_dir / "B" / "Report.docx")["table"].parent, paths.output_dir / "B")

    def test_block_ids_are_unique_and_human_location_does_not_invent_page(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            source = paths.input_dir / "A" / "Report.docx"
            make_docx(source)

            block_map = build_block_map(source, paths.input_dir)
            block_ids = [block["block_id"] for block in block_map["blocks"]]

            self.assertEqual(len(block_ids), len(set(block_ids)))
            location = build_human_location(block_map["blocks"][0], {"page": None})
            self.assertIn("абзац", location)
            self.assertNotIn("Страница", location)

    def test_report_and_annotation_from_fake_issue(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            source = paths.input_dir / "A" / "Report.docx"
            make_docx(source)
            block_map = build_block_map(source, paths.input_dir)
            block = block_map["blocks"][0]
            outputs = artifact_paths(paths, source)
            issue = {
                "issue_id": "E001",
                "severity": "medium",
                "rule_id": "fake",
                "human_location": "Страница 1, абзац 1",
                "page": 1,
                "object_type": block["object_type"],
                "quote": "Hello",
                "problem": "Fake problem",
                "recommendation": "Fake recommendation",
                "fix_mode": "requires_review",
                "confidence": 0.9,
                "status": "requires_review",
                "block_id": block["block_id"],
                "technical_location": block["technical_location"],
                "page_source": "pdf_marker",
                "page_confidence": "high",
            }

            write_audit_table(outputs["table"], [issue])
            write_annotated_document(source, outputs["annotated"], block_map, [issue], outputs["annotation_log"])

            wb = load_workbook(outputs["table"])
            headers = [cell.value for cell in wb.active[1]]
            self.assertIn("Human Location", headers)
            self.assertTrue(outputs["annotated"].exists())
            self.assertIn("E001", outputs["annotation_log"].read_text(encoding="utf-8"))

            clean_path = outputs["annotated"].with_name("Report__unanchored.docx")
            removed = strip_error_anchors_from_document(outputs["annotated"], clean_path)
            clean_doc = Document(str(clean_path))
            clean_text = "\n".join(paragraph.text for paragraph in clean_doc.paragraphs)

            self.assertEqual(removed, 1)
            self.assertNotIn("⟦E001⟧", clean_text)

    def test_official_nested_quote_false_positive_is_suppressed(self) -> None:
        block_map = {
            "blocks": [
                {
                    "block_id": "docx_tbl_0001_r002_c002_p002",
                    "object_type": "table_cell",
                    "text": (
                        "«О государственном природном заказнике регионального значения "
                        "«Дубынский» в Казанском районе» (вместе с «Положением о "
                        "государственном природном заказнике регионального значения «Дубынский»)"
                    ),
                    "technical_location": {},
                }
            ]
        }
        raw_issues = [
            {
                "block_id": "docx_tbl_0001_r002_c002_p002",
                "quote": (
                    "«О государственном природном заказнике регионального значения "
                    "«Дубынский» в Казанском районе» (вместе с «Положением о "
                    "государственном природном заказнике регионального значения «Дубынский»"
                ),
                "violation": "Сломаны кавычки в названии документа.",
                "fix": "Добавить закрывающие кавычки в названии.",
            }
        ]

        self.assertEqual(normalize_issues(raw_issues, block_map, {"entries": []}), [])

    def test_plain_broken_quote_issue_is_kept(self) -> None:
        block_map = {
            "blocks": [
                {
                    "block_id": "docx_p_0001",
                    "object_type": "paragraph",
                    "text": "Заголовок «без закрывающей кавычки",
                    "technical_location": {},
                }
            ]
        }
        raw_issues = [
            {
                "block_id": "docx_p_0001",
                "quote": "Заголовок «без закрывающей кавычки",
                "violation": "Сломаны кавычки в названии.",
                "fix": "Добавить закрывающую кавычку.",
            }
        ]

        issues = normalize_issues(raw_issues, block_map, {"entries": []})

        self.assertEqual(len(issues), 1)
        self.assertEqual(issues[0]["problem"], "Сломаны кавычки в названии.")

    def test_russian_report_headers(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "audit_ru.xlsx"
            write_audit_table(out_path, [], report_lang="ru")
            wb = load_workbook(out_path)
            headers = [cell.value for cell in wb.active[1]]

            self.assertEqual(wb.active.title, "Аудит")
            self.assertIn("Человеческая локация", headers)
            self.assertIn("Техническая локация", headers)
            self.assertIn("Режим правки", headers)
            self.assertIn("Уверенность", headers)
            self.assertIn("Старый текст", headers)
            self.assertIn("Новый текст", headers)
            self.assertIn("ID блока", headers)
            self.assertNotIn("Критичность", headers)
            self.assertNotIn("ID правила", headers)
            self.assertNotIn("Режим исправления", headers)
            self.assertNotIn("Уверенность страницы", headers)
            self.assertNotIn("Источник страницы", headers)
            self.assertNotIn("Статус", headers)

    def test_russian_report_localizes_visible_enums(self) -> None:
        issue = {
            "issue_id": "E001",
            "object_type": "paragraph",
            "problem": "CHECK: Ошибка оформления.",
            "recommendation": "Исправить оформление.",
            "fix_mode": "requires_review",
            "confidence": "high",
            "technical_location": {"part": "word/document.xml", "paragraph_index": 1},
        }
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "audit_ru.xlsx"
            write_audit_table(out_path, [issue], report_lang="ru")
            row = [cell.value for cell in load_workbook(out_path).active[2]]

        self.assertEqual(row[3], "абзац")
        self.assertEqual(row[9], "ПРОВЕРКА: Ошибка оформления.")
        self.assertEqual(row[11], "требуется проверка")
        self.assertEqual(row[12], "высокая")
        self.assertEqual(row[16], "часть: word/document.xml; абзац: 1")

    def test_russian_prompt_and_publication_reject_english_only_explanations(self) -> None:
        instructions = build_instructions("Правила", report_lang="ru")
        self.assertIn("Write every human-readable value in Russian", instructions)
        with self.assertRaisesRegex(RuntimeError, "English-only"):
            validate_human_report_language(
                [{"issue_id": "E001", "problem": "Missing space after comma", "recommendation": "Insert space"}],
                "ru",
            )

    def test_russian_language_check_rejects_english_dominant_mixed_text(self) -> None:
        invalid = invalid_human_report_fields(
            [
                {
                    "issue_id": "E001",
                    "problem": "Remove extra space after dash, затем исправить.",
                    "recommendation": "Удалить лишний пробел после тире.",
                }
            ],
            "ru",
        )

        self.assertEqual([(item["issue_id"], item["field"]) for item in invalid], [("E001", "problem")])

    def test_russian_language_check_keeps_quoted_and_technical_latin(self) -> None:
        invalid = invalid_human_report_fields(
            [
                {
                    "issue_id": "E091",
                    "quote": "WordClass",
                    "problem": "В тексте написано WordClass вместо World Class.",
                    "recommendation": "Заменить WordClass на World Class.",
                },
                {
                    "issue_id": "E092",
                    "quote": "",
                    "problem": "Ошибка в part word/document.xml, тег w:tbl не закрыт.",
                    "recommendation": "Указать «Times New Roman» в стиле абзаца.",
                },
            ],
            "ru",
        )

        self.assertEqual(invalid, [])

    def test_language_repair_provider_chains(self) -> None:
        self.assertEqual(
            language_repair_provider_chain("openai", "gpt-5.4-mini"),
            [
                {"provider": "openai", "model": "gpt-5.4-mini"},
                {"provider": "gemini", "model": "gemini-3.6-flash"},
            ],
        )
        self.assertEqual(language_repair_provider_chain("xai", "grok-4.3")[1]["provider"], "gemini")
        self.assertEqual(language_repair_provider_chain("gemini", "gemini-3.6-flash")[1], {
            "provider": "openai",
            "model": "gpt-5.6-luna",
        })

    def test_language_repair_updates_only_invalid_fields_and_reuses_cache(self) -> None:
        issues = [
            {
                "issue_id": "E015",
                "quote": "пример",
                "problem": "Remove extra space after dash",
                "recommendation": "Убрать лишний пробел после тире.",
            }
        ]
        response = ({"rows": [{"target_id": "L001", "text": "Лишний пробел после тире."}]}, {}, "default", "low")

        with tempfile.TemporaryDirectory() as tmp:
            diagnostic = Path(tmp) / "language_repair.json"
            with patch("pipeline._call_language_repair_provider", return_value=response) as provider_call:
                repaired, meta = repair_human_report_language(
                    issues,
                    report_lang="ru",
                    provider="openai",
                    model="gpt-5.4-mini",
                    output_path=diagnostic,
                    max_output_tokens=12000,
                    timeout_sec=60,
                    max_retries=1,
                    service_tier="default",
                )
            provider_call.assert_called_once()
            self.assertIn("JSON", provider_call.call_args.kwargs["user_prompt"])
            self.assertEqual(repaired[0]["problem"], "Лишний пробел после тире.")
            self.assertEqual(repaired[0]["recommendation"], issues[0]["recommendation"])
            self.assertFalse(meta["cache_hit"])
            self.assertEqual(json.loads(diagnostic.read_text(encoding="utf-8"))["status"], "repaired")

            with patch("pipeline._call_language_repair_provider") as cached_provider_call:
                cached, cached_meta = repair_human_report_language(
                    issues,
                    report_lang="ru",
                    provider="openai",
                    model="gpt-5.4-mini",
                    output_path=diagnostic,
                    max_output_tokens=12000,
                    timeout_sec=60,
                    max_retries=1,
                    service_tier="default",
                )
            cached_provider_call.assert_not_called()
            self.assertEqual(cached[0]["problem"], "Лишний пробел после тире.")
            self.assertTrue(cached_meta["cache_hit"])

    def test_language_repair_rejects_second_language_failure_and_keeps_diagnostic(self) -> None:
        issues = [
            {
                "issue_id": "E016",
                "quote": "пример",
                "problem": "Insert a space after comma",
                "recommendation": "Добавить пробел после запятой.",
            }
        ]
        response = ({"rows": [{"target_id": "L001", "text": "Still written in English"}]}, {}, "default", "low")

        with tempfile.TemporaryDirectory() as tmp:
            diagnostic = Path(tmp) / "language_repair.json"
            with patch("pipeline._call_language_repair_provider", return_value=response) as provider_call:
                with self.assertRaisesRegex(RuntimeError, "no audit report was published"):
                    repair_human_report_language(
                        issues,
                        report_lang="ru",
                        provider="openai",
                        model="gpt-5.4-mini",
                        output_path=diagnostic,
                        max_output_tokens=12000,
                        timeout_sec=60,
                        max_retries=1,
                        service_tier="default",
                    )

            self.assertEqual(provider_call.call_count, 2)
            payload = json.loads(diagnostic.read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "failed")
            self.assertIn("English-only", payload["error"])
            self.assertEqual([item["provider"] for item in payload["attempts"]], ["openai", "gemini"])

    def test_language_repair_uses_second_provider_after_invalid_first_repair(self) -> None:
        issues = [
            {
                "issue_id": "E017",
                "quote": "пример",
                "problem": "Remove spaces around dash",
                "recommendation": "Убрать пробелы вокруг тире.",
            }
        ]
        invalid = ({"rows": [{"target_id": "L001", "text": "Still written in English"}]}, {}, "default", "low")
        valid = ({"rows": [{"target_id": "L001", "text": "Лишние пробелы вокруг тире."}]}, {}, "default", "minimal")

        with tempfile.TemporaryDirectory() as tmp:
            diagnostic = Path(tmp) / "language_repair.json"
            with patch("pipeline._call_language_repair_provider", side_effect=[invalid, valid]) as provider_call:
                repaired, meta = repair_human_report_language(
                    issues,
                    report_lang="ru",
                    provider="openai",
                    model="gpt-5.4-mini",
                    output_path=diagnostic,
                    max_output_tokens=12000,
                    timeout_sec=60,
                    max_retries=1,
                    service_tier="default",
                )

            self.assertEqual(provider_call.call_count, 2)
            self.assertEqual(provider_call.call_args_list[0].kwargs["provider"], "openai")
            self.assertEqual(provider_call.call_args_list[1].kwargs["provider"], "gemini")
            self.assertEqual(provider_call.call_args_list[1].kwargs["model"], "gemini-3.6-flash")
            self.assertEqual(repaired[0]["problem"], "Лишние пробелы вокруг тире.")
            self.assertTrue(meta["fallback_used"])
            self.assertEqual(meta["provider"], "gemini")

            payload = json.loads(diagnostic.read_text(encoding="utf-8"))
            self.assertEqual(payload["status"], "repaired")
            self.assertTrue(payload["fallback_used"])
            self.assertEqual(payload["repair_provider"], "gemini")
            self.assertEqual(len(payload["attempts"]), 2)

    def test_doc_task_inputs_include_xlsx_without_changing_audit_scan(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            make_docx(paths.input_dir / "A" / "Report.docx")
            make_xlsx(paths.input_dir / "B" / "Objects.xlsx")

            audit_docs = [p.relative_to(paths.input_dir).as_posix() for p in iter_documents(paths.input_dir)]
            task_inputs = [p.relative_to(paths.input_dir).as_posix() for p in iter_task_inputs(paths.input_dir)]

            self.assertEqual(audit_docs, ["A/Report.docx"])
            self.assertEqual(task_inputs, ["A/Report.docx", "B/Objects.xlsx"])

            xlsx_map = build_task_block_map(paths.input_dir / "B" / "Objects.xlsx", paths.input_dir)
            self.assertEqual(xlsx_map["document_type"], "xlsx")
            self.assertIn("Central Clinic", xlsx_map["blocks"][1]["text"])

            docx_map = build_task_block_map(paths.input_dir / "A" / "Report.docx", paths.input_dir)
            corpus = build_corpus_block_map([docx_map, xlsx_map])
            self.assertEqual(corpus["document_type"], "corpus")
            self.assertTrue(all("__" in block["block_id"] for block in corpus["blocks"]))
            self.assertEqual(len(corpus["documents"]), 2)

    def test_doc_task_inputs_include_pdf_first_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            make_pdf(paths.input_dir / "Project.pdf")

            task_inputs = [p.relative_to(paths.input_dir).as_posix() for p in iter_task_inputs(paths.input_dir)]
            self.assertEqual(task_inputs, ["Project.pdf"])

            pdf_map = build_task_block_map(paths.input_dir / "Project.pdf", paths.input_dir, pdf_max_pages=1)
            self.assertEqual(pdf_map["document_type"], "pdf")
            self.assertEqual(len(pdf_map["blocks"]), 1)
            self.assertEqual(pdf_map["blocks"][0]["block_id"], "pdf_p_0001")
            self.assertIn("Project name: North Plant", pdf_map["blocks"][0]["text"])
            self.assertNotIn("second page", pdf_map["blocks"][0]["text"])

    def test_doc_task_prompt_mentions_json_for_openai_response_format(self) -> None:
        block = {
            "block_id": "pdf_p_0001",
            "object_type": "pdf_page",
            "source_relative_path": "Project.pdf",
            "location": "Страница 1",
            "text": "Project name: North Plant",
        }
        prompt = build_prompt(
            "Extract requisites.",
            "Return rows.",
            {"source_relative_path": "Project.pdf", "document_type": "pdf", "blocks": [block]},
            [],
            [block],
        )

        self.assertIn("JSON object", prompt)

    def test_doc_task_xlsx_writes_dynamic_values_columns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "doc_task.xlsx"
            write_task_xlsx(
                out_path,
                {
                    "rows": [
                        {
                            "source_relative_path": "A/Report.docx",
                            "block_id": "d001__docx_p_0001",
                            "quote": "Central Clinic",
                            "result": "matched",
                            "notes": "",
                            "values": {
                                "object_name": "Central Clinic",
                                "match_status": "matched",
                            },
                        }
                    ],
                    "replacements": [],
                    "document_items": [
                        {"source_relative_path": "A/Report.docx", "document_type": "docx", "blocks": 1}
                    ],
                },
            )

            wb = load_workbook(out_path)
            headers = [cell.value for cell in wb["Results"][1]]

            self.assertIn("object_name", headers)
            self.assertIn("match_status", headers)
            self.assertEqual(wb["Documents"]["A2"].value, "A/Report.docx")

    def test_doc_task_docx_report_is_word_readable(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "doc_task.docx"
            payload = {
                "task_label": "Match objects",
                "provider": "gemini",
                "model": "gemini-2.5-flash-lite",
                "task_scope": "corpus",
                "documents": 2,
                "summary": "Processed 2 input files.",
                "rows": [
                    {
                        "source_relative_path": "A/Report.docx",
                        "block_id": "d001__docx_p_0001",
                        "quote": "Central Clinic",
                        "result": "matched",
                        "notes": "high confidence",
                        "values": {"object_name": "Central Clinic", "match_status": "matched"},
                    }
                ],
                "replacements": [],
            }

            write_task_docx(out_path, payload)
            doc = Document(str(out_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

            self.assertIn("Audion Docs AI", text)
            self.assertGreaterEqual(len(doc.tables), 2)

    def test_doc_task_clean_table_outputs_follow_docx_template(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template_path = root / "template.docx"
            report_dir = root / "out"
            make_requisites_template_docx(template_path)

            outputs = write_clean_table_outputs(
                report_dir,
                "2026-05-14_12-00-00",
                {
                    "rows": [
                        {
                            "values": {
                                "Наименование": "Большеярковское сельское поселение",
                                "Вид градостроительной документации": "Об утверждении генерального плана",
                                "Кем утверждено": "решение Совета Большеярковского сельского поселения от 28.03.2019 №111",
                                "Дата": "28.03.2019",
                                "номер": "111",
                                "source_file": "ignored.pdf",
                            }
                        }
                    ]
                },
                template_path,
            )

            self.assertTrue(Path(outputs["json"]).exists())
            self.assertTrue(Path(outputs["xlsx"]).exists())
            self.assertTrue(Path(outputs["docx"]).exists())

            wb = load_workbook(outputs["xlsx"])
            ws = wb.active
            self.assertEqual([cell.value for cell in ws[1]], [
                "Наименование",
                "Вид градостроительной документации",
                "Кем утверждено",
                "Дата",
                "номер",
            ])
            self.assertEqual(ws["C2"].value, "решение Совета Большеярковского сельского поселения")
            self.assertEqual(ws["D2"].value, "от 28.03.2019")
            self.assertEqual(ws["E2"].value, "№ 111")
            self.assertNotIn("source_file", [cell.value for cell in ws[1]])

            doc = Document(outputs["docx"])
            self.assertEqual(len(doc.tables[0].rows), 2)
            self.assertEqual(doc.tables[0].cell(1, 2).text, "решение Совета Большеярковского сельского поселения")
            self.assertEqual(doc.tables[0].cell(1, 3).text, "от 28.03.2019")
            self.assertEqual(doc.tables[0].cell(1, 4).text, "№ 111")

    def test_doc_task_clean_template_auto_detection_requires_other_sources(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_dir = root / "input"
            template_path = input_dir / "template.docx"
            pdf_path = input_dir / "source.pdf"
            make_requisites_template_docx(template_path)
            make_pdf(pdf_path)

            self.assertIsNone(resolve_clean_template(input_dir, [template_path], ""))
            self.assertEqual(resolve_clean_template(input_dir, [template_path, pdf_path], ""), template_path.resolve())

    def test_llm_cache_signature_rejects_same_name_edited_source(self) -> None:
        sig = llm_cache_signature(
            source_sha256="old",
            provider="openai",
            model="gpt-test",
            reasoning="low",
            rules_context="rules",
            instructions="instructions",
            chunk_tokens=1000,
            overlap_tokens=100,
            min_chunks=1,
            max_output_tokens=2000,
        )
        edited = dict(sig)
        edited["source_sha256"] = "new"

        self.assertEqual(llm_cache_mismatch_reason({"cache_signature": sig}, sig), "")
        self.assertEqual(llm_cache_mismatch_reason({"cache_signature": sig}, edited), "source_sha256 changed")
        self.assertEqual(llm_cache_mismatch_reason({"rows": []}, sig), "missing cache_signature")

    def test_doc_task_cache_uses_content_hash_not_only_filename(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            paths = default_project_paths(Path(tmp))
            ensure_project_dirs(paths)
            source = paths.input_dir / "Report.docx"
            make_docx(source, "First version")

            first = source_inventory_hash([source], paths.input_dir, "task")
            make_docx(source, "Edited version")
            second = source_inventory_hash([source], paths.input_dir, "task")

            self.assertNotEqual(first, second)

            sig = doc_task_cache_signature(
                source_sha256="old",
                task_ref="active",
                task_instruction="Extract values",
                user_query="Return rows",
                provider="openai",
                model="gpt-test",
                chunk_tokens=1000,
                overlap_tokens=0,
                min_chunks=1,
                max_output_tokens=2000,
                pdf_max_pages=5,
                system_prompt_sha256="system",
            )
            edited = dict(sig)
            edited["source_sha256"] = "new"

            self.assertEqual(doc_task_cache_mismatch_reason({"cache_signature": sig}, sig), "")
            self.assertEqual(doc_task_cache_mismatch_reason({"cache_signature": sig}, edited), "source_sha256 changed")
            self.assertEqual(doc_task_cache_mismatch_reason({"rows": []}, sig), "missing cache_signature")

    def test_doc_task_safe_replacements_leave_unresolved_report_items(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source.docx"
            make_docx(source, "Fix teh word and keep MaybeName.")
            out_docx = root / "source__normalized.docx"
            replacements = [
                {
                    "replacement_id": "X001",
                    "source_relative_path": "source.docx",
                    "block_id": "docx_p_0001",
                    "old_text": "teh",
                    "new_text": "the",
                    "reason": "obvious typo",
                    "confidence": "high",
                },
                {
                    "replacement_id": "X002",
                    "source_relative_path": "source.docx",
                    "block_id": "docx_p_0001",
                    "old_text": "MaybeName",
                    "new_text": "Maybe Name",
                    "reason": "could be a proper name",
                    "confidence": "medium",
                },
                {
                    "replacement_id": "X003",
                    "source_relative_path": "source.docx",
                    "block_id": "docx_p_0001",
                    "old_text": "missing text",
                    "new_text": "replacement",
                    "reason": "not present",
                    "confidence": "high",
                },
            ]

            edit = apply_docx_replacements_safely(
                source,
                out_docx,
                replacements,
                min_confidence="high",
                create_copy=True,
            )

            fixed = Document(str(out_docx))
            fixed_text = "\n".join(paragraph.text for paragraph in fixed.paragraphs)
            self.assertEqual(edit["applied"], 1)
            self.assertIn("Fix the word", fixed_text)
            self.assertIn("MaybeName", fixed_text)
            self.assertEqual(replacements[0]["status"], "applied")
            self.assertEqual(replacements[1]["status"], "unresolved")
            self.assertEqual(replacements[2]["status"], "unresolved")

            unresolved = build_unresolved_items([], replacements)
            self.assertEqual(len(unresolved), 2)
            self.assertIn("confidence medium", unresolved[0]["reason"])
            self.assertIn("old_text", unresolved[1]["reason"])

            report_docx = root / "report.docx"
            payload = {
                "task_label": "Safe replacements",
                "provider": "openai",
                "model": "test-model",
                "task_scope": "document",
                "documents": 1,
                "summary": "Applied one safe replacement.",
                "rows": [],
                "replacements": replacements,
                "unresolved_items": unresolved,
            }
            write_task_docx(report_docx, payload)
            report = Document(str(report_docx))
            report_text = "\n".join(paragraph.text for paragraph in report.paragraphs)
            self.assertIn("Unresolved Items", report_text)
            self.assertTrue(any("MaybeName" in cell.text for table in report.tables for row in table.rows for cell in row.cells))

    def test_document_normalizer_uses_existing_openai_audit_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_dir = root / "input"
            logs_dir = root / "logs"
            output_dir = root / "output"
            report_dir = output_dir / "_normalization"
            patch_dir = root / "report" / "document_normalization"
            source = input_dir / "source.docx"
            make_docx(source, "Fix teh word and keep MaybeName.")
            audit_path = logs_dir / "source__audit.json"
            audit_path.parent.mkdir(parents=True, exist_ok=True)
            audit_path.write_text(
                json.dumps(
                    {
                        "source_relative_path": "source.docx",
                        "status": "ok",
                        "meta": {"provider": "openai"},
                        "issues": [
                            {
                                "issue_id": "E001",
                                "block_id": "docx_p_0001",
                                "quote": "teh",
                                "problem": "typo",
                                "recommendation": "replace with the",
                                "fix_mode": "safe_replace",
                                "old_text": "teh",
                                "new_text": "the",
                                "confidence": "high",
                            },
                            {
                                "issue_id": "E002",
                                "block_id": "docx_p_0001",
                                "quote": "MaybeName",
                                "problem": "possible spacing",
                                "recommendation": "manual review",
                                "fix_mode": "requires_review",
                                "confidence": "medium",
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            payload = normalize_from_audit_logs(
                root=root,
                provider="openai",
                logs_dir=logs_dir,
                input_dir=input_dir,
                output_dir=output_dir,
                report_dir=report_dir,
                patch_dir=patch_dir,
            )

            normalized = output_dir / "source__normalized.docx"
            fixed = Document(str(normalized))
            fixed_text = "\n".join(paragraph.text for paragraph in fixed.paragraphs)
            self.assertIn("Fix the word", fixed_text)
            self.assertIn("MaybeName", fixed_text)
            self.assertEqual(len(payload["applied_items"]), 1)
            self.assertEqual(len(payload["unresolved_items"]), 1)
            self.assertTrue(list(report_dir.glob("*__normalization_openai.docx")))
            self.assertTrue((patch_dir / "latest_normalization_patch_plan_openai.json").exists())

    def test_quick_doc_task_env_text_overrides_file_task(self) -> None:
        old_value = os.environ.get(TASK_TEXT_ENV)
        try:
            os.environ[TASK_TEXT_ENV] = "Extract names into values.object_name."
            instruction, entry = load_doc_task("")

            self.assertIn("Extract names", instruction)
            self.assertTrue(entry.get("quick"))
            self.assertTrue(str(entry.get("ref", "")).startswith("quick:"))
        finally:
            if old_value is None:
                os.environ.pop(TASK_TEXT_ENV, None)
            else:
                os.environ[TASK_TEXT_ENV] = old_value

    def test_quick_doc_task_delete_removes_cache_entry_and_pin(self) -> None:
        old_cache_path = doc_task_resolver.QUICK_TASK_CACHE_PATH
        with tempfile.TemporaryDirectory() as tmp:
            try:
                doc_task_resolver.QUICK_TASK_CACHE_PATH = Path(tmp) / "quick_cache.json"
                entry = doc_task_resolver.save_quick_doc_task(
                    "Match DOCX rows to XLSX rows.",
                    label="Temporary quick task",
                    pin=True,
                )

                deleted = doc_task_resolver.delete_quick_doc_task(str(entry["ref"]))
                entries = doc_task_resolver.quick_doc_task_entries()

                self.assertEqual(deleted["label"], "Temporary quick task")
                self.assertEqual(entries, [])
                self.assertEqual(doc_task_resolver.pinned_quick_doc_task_refs(), [])
            finally:
                doc_task_resolver.QUICK_TASK_CACHE_PATH = old_cache_path


if __name__ == "__main__":
    unittest.main()
